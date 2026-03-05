# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("INFORME_FINAL_SISTEMA.docx")

textos = {
    "UML: Diagrama de Componentes": "El Diagrama de Componentes UML representa las unidades de software del Sistema de Gestión de Tareas  Oficina EPO y sus dependencias. Incluye el frontend React (componentes, páginas, contextos y utilidades), el backend Node.js/Express (rutas, middleware, modelos Mongoose) y la base de datos MongoDB Atlas. Las interfaces entre componentes están definidas mediante la API RESTful con endpoints organizados por recursos (auth, tasks, reports, notifications).",
    "UML: Diagrama de Casos de Uso": "El Diagrama de Casos de Uso describe las interacciones entre los actores del sistema y sus funcionalidades. Los actores son: Jefa del CMC (rol administrador) y Asistente del CMC (rol usuario). Los casos de uso principales son: CS-01 Visualizar Dashboard, CS-02 Gestionar Tareas, CS-03 Generar Reporte PDF, CS-04 Gestionar Notificaciones, CS-05 Ver Estadísticas, CS-06 Autenticar Usuario y CS-07 Gestionar Perfil de Usuario.",
    "Diagramas de Procesos": "Los Diagramas de Secuencia UML documentan el flujo temporal de mensajes entre los actores y el sistema. Se modelan los procesos principales: inicio de sesión con autenticación JWT, creación y actualización de tareas con cambio de estado Kanban, generación automática de reportes PDF con jsPDF, y el flujo de notificaciones automáticas para tareas vencidas. Evidencian la separación de responsabilidades entre el cliente React y el servidor Node.js/Express conectado a MongoDB Atlas.",
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for titulo, contenido in textos.items():
        if text == titulo:
            # Ver si el siguiente párrafo tiene contenido relevante
            siguiente_text = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ""
            if len(siguiente_text) < 30 or siguiente_text.startswith("UML:") or siguiente_text.startswith("Diagrama") or siguiente_text.startswith("Tecnologías"):
                # Insertar nuevo párrafo inmediatamente después vía XML
                new_para_xml = OxmlElement("w:p")
                new_run = OxmlElement("w:r")
                new_text = OxmlElement("w:t")
                new_text.text = contenido
                new_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                new_run.append(new_text)
                new_para_xml.append(new_run)
                p._p.addnext(new_para_xml)
                print(f"  Restaurado: {titulo}")
            break

doc.save("INFORME_FINAL_SISTEMA.docx")
print("Listo.")

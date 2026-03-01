import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)
p    = doc.paragraphs

def ft(para):
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))

def set_xml(para, new_text):
    all_t = para._p.findall(f'.//{{{NS}}}t')
    if all_t:
        all_t[0].text = new_text
        for t in all_t[1:]: t.text = ''
    else:
        para.add_run(new_text)

# ── [3757] Diagrama de despliegue completo ─────────────────────────────────
# Reemplazar con descripción de la arquitectura real del sistema
nueva_arq = (
    "Cliente (Navegador Web) — Chrome / Firefox / Edge\n"
    "HTTPS | Puerto 443\n"
    "Frontend: Vercel (plan gratuito)\n"
    "  React 18 + Vite 5 | Build: npm run build | Auto-deploy desde GitHub\n"
    "  CDN edge global | SSL/TLS automático\n"
    "HTTPS/REST API → https://api.render.com | Puerto 443\n"
    "Backend: Render (plan gratuito)\n"
    "  Node.js 18.x + Express 4.x\n"
    "  Proceso: web service | Puerto: 10000 (Render asigna)\n"
    "  Variables de entorno: .env (MONGO_URI, JWT_SECRET)\n"
    "  API Routes: /api/auth | /api/tareas | /api/notificaciones | /api/reportes\n"
    "MONGODB+SRV (TLS) → Puerto 27017\n"
    "Base de datos: MongoDB Atlas (plan gratuito M0)\n"
    "  Región: US East | Replica Set: 3 nodos | Auto-failover\n"
    "  Respaldos automatizados | Colecciones: usuarios, tareas, notificaciones\n"
    "  Mongoose ODM | Connection Pool\n"
    "  Autenticación: JWT (Bearer Token) | Hashing: bcrypt (cost 10)"
)
set_xml(p[3757], nueva_arq)
print(f'[3757] reemplazado ({len(nueva_arq)} chars)')

# Verificar que no quede nginx
txt = ft(p[3757]).lower()
if 'nginx' in txt:
    print('  ⚠ nginx aún presente')
else:
    print('  ✓ nginx eliminado')

# ── Verificar otros párrafos cercanos que puedan tener residuales ──────────
bad_terms = ['nginx', 'ubuntu server', 'aws s3', 'gmail/sendgrid', 'oauth 2', 'sendgrid', 'pm2', 'ufw']
for i in range(3750, 3770):
    t = ft(p[i]).lower()
    for b in bad_terms:
        if b in t:
            print(f'  [{i}] tiene "{b}": {ft(p[i])[:80]}')

doc.save(DEST)
print('\n✅ Guardado.')

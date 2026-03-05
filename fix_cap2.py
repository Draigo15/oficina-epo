# -*- coding: utf-8 -*-
from docx import Document
from copy import deepcopy
from lxml import etree

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"

idx_xp = idx_mongo = idx_router = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == "Extreme Programming (XP)" and p.style.name.startswith("Heading"):
        idx_xp = i
    if t == "Base de Datos NoSQL: MongoDB" and p.style.name.startswith("Heading"):
        idx_mongo = i
    if t == "Enrutamiento: React Router DOM" and p.style.name.startswith("Heading"):
        idx_router = i

print(f"XP: {idx_xp}, MongoDB: {idx_mongo}, Router: {idx_router}")

# Buscar template Body Text
tmpl = None
for p in paras:
    if p.style.name == "Body Text" and p.text.strip():
        tmpl = p
        break

def insert_body_after(idx, text):
    ref = paras[idx]._p
    new_p = deepcopy(tmpl._p)
    # Limpiar runs
    for r in new_p.findall(f"{{{NS}}}r"):
        new_p.remove(r)
    # Crear run
    run_el = etree.SubElement(new_p, f"{{{NS}}}r")
    tmpl_runs = tmpl._p.findall(f"{{{NS}}}r")
    if tmpl_runs:
        rpr = tmpl_runs[0].find(f"{{{NS}}}rPr")
        if rpr is not None:
            run_el.insert(0, deepcopy(rpr))
    t_el = etree.SubElement(run_el, f"{{{NS}}}t")
    t_el.set(f"{{{XML_NS}}}space", "preserve")
    t_el.text = text
    ref.addnext(new_p)

txt_xp = (
    "Extreme Programming (XP) es una metodología ágil de desarrollo de software que prioriza la entrega "
    "continua de software funcional en iteraciones cortas. En el Sistema de Gestión de Tareas \u2013 Oficina EPO "
    "se adoptó XP como marco de trabajo, aplicando prácticas como entregas incrementales, integración continua "
    "y retroalimentación constante del usuario final. Cada iteración incorporó funcionalidades completas "
    "(gestión de tareas, autenticación, generación de reportes PDF y notificaciones), permitiendo validar "
    "el producto de forma temprana y ajustar los requerimientos de manera flexible durante el desarrollo."
)

txt_mongo = (
    "MongoDB es un sistema de base de datos NoSQL orientado a documentos que almacena la información en formato "
    "BSON (Binary JSON). En el Sistema de Gestión de Tareas \u2013 Oficina EPO, MongoDB Atlas se utiliza como servicio "
    "en la nube para persistir tres colecciones principales: Users (datos de usuarios y credenciales hasheadas con "
    "bcryptjs), Tasks (tareas con título, descripción, estado, prioridad, fecha límite y usuario asignado) y "
    "Notifications (alertas generadas automáticamente al crear, actualizar o eliminar tareas). La conexión se "
    "gestiona mediante Mongoose y la URI se almacena de forma segura en variables de entorno."
)

txt_router = (
    "React Router DOM es la biblioteca estándar de enrutamiento para aplicaciones React en el navegador. "
    "En el Sistema de Gestión de Tareas \u2013 Oficina EPO se utiliza para gestionar la navegación entre las "
    "distintas vistas de la aplicación: Dashboard, Tareas, Reportes, Notificaciones, Estadísticas y Perfil. "
    "Se implementan rutas protegidas mediante un componente PrivateRoute que verifica la autenticación JWT "
    "antes de permitir el acceso, redirigiendo al Login cuando el token no es válido o ha expirado."
)

# Insertar de abajo hacia arriba
for idx, txt in sorted([(idx_router, txt_router), (idx_mongo, txt_mongo), (idx_xp, txt_xp)], key=lambda x: x[0], reverse=True):
    insert_body_after(idx, txt)
    print(f"  Texto insertado despues de parrafo {idx}")

# Eliminar parrafo huerfano Recharts (buscar por texto unico)
for i, p in enumerate(doc.paragraphs):
    if "Recharts es una biblioteca de gráficos basada en SVG para React que proporciona componentes declarativos" in p.text:
        p._p.getparent().remove(p._p)
        print(f"  Parrafo huerfano Recharts eliminado (indice {i})")
        break

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"\nGuardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"\nArchivo bloqueado, guardado en {out}")

# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

# === 1. Encontrar inicio de los RF del portal (RF-004) y fin (antes de RNF) ===
inicio_portal_rf = None
fin_portal_rf = None

for i, p in enumerate(paras):
    t = p.text.strip()
    # Inicio: el heading "Modulo de Prospectos" o el parrafo de RF-004 con formularios
    if "Módulo de Prospectos" in t or "formularios p\u00FAblicos" in t.lower() or (
        "RF-" in t and "formularios" in paras[i+2].text.lower() if i+2 < len(paras) else False):
        if inicio_portal_rf is None and i > 900:
            inicio_portal_rf = i
            print(f"Inicio portal RF: [{i}] '{t[:60]}'")

for i, p in enumerate(paras):
    t = p.text.strip()
    if "Requerimientos No Funcionales" in t and p.style.name == "Heading 3" and i > 900:
        fin_portal_rf = i - 1
        print(f"Fin portal RF: [{fin_portal_rf}] '{paras[fin_portal_rf].text[:60]}'")
        break

if inicio_portal_rf is None:
    # Buscar por RF-004 con contenido de leads
    for i, p in enumerate(paras):
        if "RF-" in p.text and i > 920:
            ctx = " ".join(paras[j].text for j in range(i, min(i+5, len(paras))))
            if "formulario" in ctx.lower() or "captaci\u00F3n" in ctx.lower() or "lead" in ctx.lower():
                inicio_portal_rf = i
                print(f"Inicio portal RF (fallback): [{i}] '{p.text[:60]}'")
                break

print(f"Rango a reemplazar: [{inicio_portal_rf}..{fin_portal_rf}]")

# === 2. Eliminar parrafos del portal RF ===
if inicio_portal_rf and fin_portal_rf:
    count = 0
    for idx in range(fin_portal_rf, inicio_portal_rf - 1, -1):
        p = doc.paragraphs[idx]
        p._p.getparent().remove(p._p)
        count += 1
    print(f"Eliminados: {count} parrafos RF portal")

# === 3. Insertar RF correctos del sistema TareasEpo ===
# El ancla ahora es "Requerimientos No Funcionales"
anchor = None
for i, p in enumerate(doc.paragraphs):
    if "Requerimientos No Funcionales" in p.text and p.style.name == "Heading 3" and i > 700:
        anchor = p
        print(f"Ancla RNF en [{i}]")
        break

if anchor is None:
    print("ERROR: No se encontro RNF como ancla")
else:
    def ins(text, style):
        p = doc.add_paragraph(style=style)
        if text:
            p.add_run(text)
        anchor._p.addprevious(p._p)

    def blank():
        ins("", "Normal")

    # Modulo de Gestion de Tareas
    ins("Módulo de Gestión de Tareas", "Heading 3")
    ins("Cód     Requerimiento   Descripción     Prioridad Justificación", "Heading 3")
    ins("RF-004  Gestionar tareas (CRUD)  Crear, editar, cambiar estado (pendiente/en progreso/completada) "
        "y eliminar tareas; asignar responsable, prioridad y fecha límite.  "
        "Alta    Funcionalidad central del sistema; requerida para operación diaria.", "Body Text")
    ins("RF-005  Cambio de estado Drag & Drop  Mover tarjetas entre columnas del tablero Kanban "
        "mediante arrastre para actualizar el estado de la tarea.  "
        "Alta    Mejora la experiencia sin necesidad de abrir formularios.", "Body Text")
    ins("RF-006  Filtros y búsqueda  Filtrar tareas por estado, prioridad, responsable y búsqueda textual.  "
        "Media   Optimiza la gestión cuando hay muchas tareas registradas.", "Body Text")

    blank()
    ins("Módulo de Reportes PDF", "Heading 3")
    ins("Cód     Requerimiento   Descripción     Prioridad Justificación", "Heading 3")
    ins("RF-007  Generar reporte mensual PDF  Seleccionar periodo (mes/año), generar PDF automático "
        "con tabla de tareas, estadísticas y KPIs mediante jsPDF y jsPDF-autotable.  "
        "Alta    Sustitución del proceso manual de consolidación; ahorro de horas.", "Body Text")
    ins("RF-008  Descarga directa  El PDF generado se descarga automáticamente en el navegador del usuario.  "
        "Alta    Proceso completamente cliente-lado, sin carga adicional al servidor.", "Body Text")

    blank()
    ins("Módulo de Notificaciones", "Heading 3")
    ins("Cód     Requerimiento   Descripción     Prioridad Justificación", "Heading 3")
    ins("RF-009  Notificaciones automáticas  Enviar notificación interna al responsable asignado "
        "al crear una nueva tarea.  "
        "Alta    Comunicación inmediata dentro del equipo sin herramientas externas.", "Body Text")
    ins("RF-010  Centro de notificaciones  Listar notificaciones del usuario autenticado, "
        "marcar como leídas, eliminar; mostrar badge con conteo de no leídas.  "
        "Media   Trazabilidad de comunicaciones internas del área.", "Body Text")
    ins("RF-011  Notificaciones manuales  La Jefa puede crear notificaciones personalizadas "
        "para el equipo desde el panel de administración.  "
        "Media   Comunicación interna directa sin depender de email.", "Body Text")

    blank()
    ins("Módulo de Dashboard y Estadísticas", "Heading 3")
    ins("Cód     Requerimiento   Descripción     Prioridad Justificación", "Heading 3")
    ins("RF-012  Dashboard principal  Mostrar KPI cards con totales de tareas por estado; "
        "gráfico de barras de completadas por semana y dona de distribución por prioridad.  "
        "Alta    Visibilidad operativa inmediata del estado del área.", "Body Text")
    ins("RF-013  Estadísticas avanzadas  Panel de estadísticas con gráficos de líneas, barras "
        "apiladas y tabla KPI: tasa de completitud, tiempo promedio, responsable más productivo.  "
        "Alta    Apoyo a decisiones de la Jefa del CMC.", "Body Text")

    blank()
    ins("Módulo de Perfil y Gestión de Usuarios", "Heading 3")
    ins("Cód     Requerimiento   Descripción     Prioridad Justificación", "Heading 3")
    ins("RF-014  Perfil de usuario  Cada usuario puede ver y actualizar sus datos personales "
        "(nombre, email) y cambiar su contraseña desde el módulo de perfil.  "
        "Media   Autonomía del usuario para mantener sus datos actualizados.", "Body Text")
    ins("RF-015  Gestión de cuentas (Jefa)  La Jefa puede crear cuentas, asignar roles "
        "(Jefa/Asistente) y activar/desactivar usuarios desde el panel de administración.  "
        "Alta    Control centralizado del acceso al sistema.", "Body Text")
    blank()

    print("RF TareasEpo insertados correctamente.")

# === 4. Limpiar diagrama BD portal (texto espaciado) en zona ~[1192-1260] ===
# Buscar patron espaciado tipico del diagrama BD portal
import re
def es_espaciado_bd(text):
    matches = re.findall(r'\b[A-Za-z\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1]{1,2}\s(?=[a-z\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1])', text)
    return len(matches) >= 3

to_del = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    if es_espaciado_bd(t) and p.style.name == "Normal" and i > 900:
        to_del.append(i)

print(f"\nTexto espaciado BD a eliminar: {len(to_del)}")
for idx in to_del[:10]:
    print(f"  [{idx}] '{doc.paragraphs[idx].text[:60]}'")

for idx in sorted(to_del, reverse=True):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)

print(f"Eliminados: {len(to_del)} parrafos BD espaciados")

print(f"\nTotal parrafos final: {len(doc.paragraphs)}")
out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

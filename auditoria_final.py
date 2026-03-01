import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx')

bad = [
    'multer', 'cloudinary', 'nodemailer', 'docker compose', 'nginx', 'openapi 3',
    'rhf + zod', 'react hook form', 'siga odontolog', '16 semanas',
    '2,720', '3,020', 'vps basico', 'vps b\u00e1sico',
    'servidor vps'
]

print('=== AUDITORÍA FINAL ===')
issues = []

for i, para in enumerate(doc.paragraphs):
    txt = para.text.lower()
    for b in bad:
        if b in txt:
            issues.append(f'Para [{i}]: "{b}" -> {para.text[:80]}')

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.lower()
            for b in bad:
                if b in txt:
                    issues.append(f'T{ti}R{ri}C{ci}: "{b}" -> {cell.text[:60]}')

if issues:
    print(f'PROBLEMAS ({len(issues)}):')
    for iss in issues:
        print(f'  {iss}')
else:
    print('✅ Sin residuales detectados')

print(f'\nTotal tablas: {len(doc.tables)} | Total párrafos: {len(doc.paragraphs)}')

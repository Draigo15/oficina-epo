"""
Script v3b — Limpieza profunda de todos los residuales detectados en auditoría
"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from lxml import etree

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc = Document(DEST)
paras = doc.paragraphs
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def set_para(para, text):
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ''
    if not para.runs:
        para.add_run(text)

def replace_in_tables(old, new):
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old.lower()[:30] in cell.text.lower():
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1
    return count

# ─── fixar texto de todos los w:t de un párrafo (incluyendo hyperlinks) ───
def fix_all_t(para_elem, old, new):
    count = 0
    for t in para_elem.findall(f'.//{{{NS}}}t'):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            count += 1
    return count

# ══════════════════════════════════════════════════════════════════════════════
# 1. ÍNDICE DE FIGURAS: entradas del TOC con "SIGA Odontología"
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- ÍNDICE / CAPTIONS SIGA ---")
for idx in [85, 86, 112]:
    n = fix_all_t(paras[idx]._p, 'SIGA Odontología', 'Sistema de Gestión de Tareas – Oficina EPO')
    n += fix_all_t(paras[idx]._p, 'SIGA Odontolog\u00eda', 'Sistema de Gestión de Tareas – Oficina EPO')
    print(f'  [{idx}] {n}x -> {paras[idx].text[:70]}')

# También buscar todos los captions de figuras en el documento
for i, para in enumerate(paras):
    if 'SIGA Odontolog' in para.text and ('Figura' in para.text or 'figura' in para.text
                                          or 'Imagen' in para.text or 'Cronograma' in para.text):
        n = fix_all_t(para._p, 'SIGA Odontología', 'Sistema Gestión Tareas – Oficina EPO')
        n += fix_all_t(para._p, 'SIGA Odontolog\u00eda', 'Sistema Gestión Tareas – Oficina EPO')
        if n:
            print(f'  [{i}] caption corregido: {para.text[:70]}')

# ══════════════════════════════════════════════════════════════════════════════
# 2. MARCO TEÓRICO — React Hook Form (Cap II)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- MARCO TEÓRICO React Hook Form ---")

# [232] React en el sistema
set_para(paras[232],
    "React es empleado para construir la interfaz de usuario del Sistema de "
    "Gestión de Tareas – Oficina EPO mediante componentes funcionales y hooks. "
    "La arquitectura de componentes reutilizables (Layout, TaskCard, Notificaciones, "
    "formularios de login y perfil) permite desarrollar una SPA (Single Page "
    "Application) performante con React Router DOM v6 para la navegación entre "
    "módulos del sistema.")
print(f"  [232] OK")

# [238] React Hook Form párrafo principal
set_para(paras[238],
    "jsPDF es una biblioteca JavaScript de código abierto que permite generar "
    "documentos PDF directamente desde el navegador, sin necesidad de servidor "
    "externo. jsPDF-autotable extiende esta funcionalidad con soporte para "
    "tablas formateadas, encabezados y pies de página, lo que facilita la "
    "exportación de reportes mensuales de tareas organizados por estado, "
    "prioridad o responsable en el Sistema de Gestión de Tareas – Oficina EPO.")
print(f"  [238] OK")

# [240] Zod párrafo
set_para(paras[240],
    "Recharts es una biblioteca de gráficos basada en SVG para React que "
    "proporciona componentes declarativos para visualizar datos estadísticos. "
    "En el sistema, Recharts se utiliza en el dashboard para mostrar gráficos "
    "de barras y líneas con las métricas de productividad del área: tareas "
    "completadas por período, distribución por prioridad y evolución histórica "
    "de la carga de trabajo. Su integración con el estado de React permite "
    "actualización en tiempo real al completar o registrar tareas.")
print(f"  [240] OK")

# ══════════════════════════════════════════════════════════════════════════════
# 3. TEST CASES con RHF+Zod (cap de pruebas)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- CASOS DE PRUEBA RHF+Zod ---")

for idx in [1093, 1218, 1256]:
    if 'react hook form' in paras[idx].text.lower() or 'zod' in paras[idx].text.lower():
        for run in paras[idx].runs:
            run.text = (run.text
                        .replace('El sistema de validación (Zod + React Hook Form) debe estar configurado.',
                                 'El sistema de autenticación JWT debe estar activo.')
                        .replace('El sistema de validación Zod + React Hook Form debe es- tar configurado.',
                                 'El sistema de autenticación JWT debe estar activo.')
                        .replace('El sistema de validación Zod + React Hook Form debe estar configurado.',
                                 'El sistema de autenticación JWT debe estar activo.')
                        .replace('El sistema renderiza formulario completo con React Hook Form.',
                                 'El sistema renderiza el formulario de tarea correctamente.')
                        .replace('El sistema renderiza formulario completo con React Hook Form',
                                 'El sistema renderiza el formulario de tarea correctamente')
                        )
        print(f'  [{idx}] -> {paras[idx].text[:80]}')

# ══════════════════════════════════════════════════════════════════════════════
# 4. TEST CASES con Multer (upload de imágenes)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- CASOS DE PRUEBA Multer ---")

for idx in [1844, 1948, 2117]:
    if 'multer' in paras[idx].text.lower():
        for run in paras[idx].runs:
            run.text = (run.text
                        .replace('El sistema de carga de imágenes (multer o MongoDB Atlas) de- be estar configura',
                                 'El sistema debe tener al menos una tarea registrada.')
                        .replace('El sistema de carga de imágenes (multer o MongoDB Atlas) debe estar configurado.',
                                 'El sistema debe tener al menos una tarea registrada.')
                        .replace('El sistema de carga de imágenes debe estar configurado (multer/MongoDB Atlas).',
                                 'El sistema debe tener al menos una tarea registrada.')
                        .replace('El backend procesa imagen con multer.',
                                 'El backend procesa la solicitud y responde con código 200.')
                        )
        print(f'  [{idx}] -> {paras[idx].text[:80]}')

# ══════════════════════════════════════════════════════════════════════════════
# 5. ÁRBOL DE ARCHIVOS con multer.middleware.js y nodemailer
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- ÁRBOL DE ARCHIVOS ---")

# [3527] multer.middleware.js
if 'multer' in paras[3527].text.lower():
    for run in paras[3527].runs:
        run.text = (run.text
                    .replace('multer.middleware.js', 'auth.middleware.js')
                    .replace('File Upload', 'Autenticación JWT')
                    )
    print(f'  [3527] -> {paras[3527].text[:70]}')

# [3699] nodemailer
if 'nodemailer' in paras[3699].text.lower():
    for run in paras[3699].runs:
        run.text = (run.text
                    .replace('nodemailer', 'date-fns')
                    .replace('Email Client', 'Date utils')
                    )
    print(f'  [3699] -> {paras[3699].text[:70]}')

# ══════════════════════════════════════════════════════════════════════════════
# 6. HARDWARE TABLE — Nginx en párrafo [3945]
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- HARDWARE Nginx ---")

if 'nginx' in paras[3945].text.lower():
    for run in paras[3945].runs:
        run.text = (run.text
                    .replace('2 vCPU dedicados (o equivalentes) para Node.js y Nginx.',
                             '2 vCPU equivalentes gestionados por Render (plan gratuito).')
                    .replace('Node.js y Nginx', 'Node.js/Express en Render')
                    )
    print(f'  [3945] -> {paras[3945].text[:80]}')

# ══════════════════════════════════════════════════════════════════════════════
# 7. TABLA SOFTWARE — residuales de Nginx, OpenAPI, RHF+Zod en párrafos
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- TABLA SOFTWARE residuales en párrafos ---")

# [4019] contiene "Tailwind CSS" + "RHF + Zod" en la misma celda de párrafo
# es un párrafo dentro de una celda de tabla, pero accesible como párrafo
if 'rhf + zod' in paras[4019].text.lower() or 'react hook form' in paras[4019].text.lower():
    for run in paras[4019].runs:
        run.text = (run.text
                    .replace('Enrutamiento (UI)    RHF + Zod     Formularios React', 
                             'Enrutamiento (UI)    React Router DOM 6.x     SPA')
                    .replace('React Hook Form', 'jsPDF-autotable')
                    .replace('RHF + Zod', 'jsPDF + auto')
                    )
    # Si no encontró por run, buscar en todos los w:t
    n = fix_all_t(paras[4019]._p, 'RHF + Zod', 'jsPDF-autotable')
    n += fix_all_t(paras[4019]._p, 'React Hook Form', 'jsPDF-autotable')
    print(f'  [4019] {n}x -> {paras[4019].text[:80]}')

# [4036] Nginx como servidor web => Git + GitHub
if 'nginx' in paras[4036].text.lower():
    n = fix_all_t(paras[4036]._p, 'Nginx (proxy)', 'Git + GitHub')
    n += fix_all_t(paras[4036]._p, 'Nginx', 'Git + GitHub')
    print(f'  [4036] {n}x -> {paras[4036].text[:80]}')

# [4039] OpenAPI
if 'openapi' in paras[4039].text.lower():
    n = fix_all_t(paras[4039]._p, 'OpenAPI 3 (opcio- nal)', 'Lucide React')
    n += fix_all_t(paras[4039]._p, 'OpenAPI 3 (opcional)', 'Lucide React')
    n += fix_all_t(paras[4039]._p, 'OpenAPI 3', 'Lucide React')
    print(f'  [4039] {n}x -> {paras[4039].text[:80]}')

# ══════════════════════════════════════════════════════════════════════════════
# 8. TABLA HARDWARE (Word table) — VPS
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- TABLA HARDWARE VPS ---")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if 'servidor vps' in cell.text.lower() or 'vps b' in cell.text.lower():
                for para in cell.paragraphs:
                    n = fix_all_t(para._p, 'Servidor VPS básico (4 meses)', 'Render (backend cloud, plan free)')
                    n += fix_all_t(para._p, 'Servidor VPS basico (4 meses)', 'Render (backend cloud, plan free)')
                    if n:
                        print(f'  Tabla celda corregida: {cell.text[:60]}')

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
doc.save(DEST)
print("\n✅ Limpieza v3b guardada.")

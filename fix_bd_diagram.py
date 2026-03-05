# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

# Verificar ancla: [1052] = "Diagrama de base de datos del sistema"
anchor_text = paras[1052].text.strip()
print(f"Ancla [1052]: '{anchor_text}'")
assert "base de datos" in anchor_text.lower(), "Ancla incorrecta!"

# Rango a eliminar: [1053..1119]  contenido portal espaciado
# El [1120] "Nota. Elaboracion propia." lo conservamos
to_delete = list(range(1053, 1120))
print(f"Eliminando {len(to_delete)} parrafos de diagrama portal...")

# Guardar referencia del parrafo ancla para insertar despues
anchor_p = paras[1052]._p

# Eliminar de atras hacia adelante
for idx in sorted(to_delete, reverse=True):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)

print("Portal BD eliminado.")
print(f"Parrafos tras eliminacion: {len(doc.paragraphs)}")

# Verificar que el siguiente parrafo del ancla es ahora "Nota. Elaboracion"
idx_after = None
for i, p in enumerate(doc.paragraphs):
    if p._p is anchor_p:
        idx_after = i
        break

if idx_after is not None:
    print(f"Ancla ahora en [{idx_after}]: '{doc.paragraphs[idx_after].text.strip()}'")
    print(f"Siguiente [{idx_after+1}]: '{doc.paragraphs[idx_after+1].text.strip()}'")

# -------------------------------------------------------
# Insertar descripcion textual del esquema real TareasEpo
# -------------------------------------------------------
def add_para_after(anchor_elem, text, style="Body Text", bold_prefix=None):
    """Inserta un parrafo despues de anchor_elem, retorna el nuevo _p."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style.replace(" ", ""))
    pPr.append(pStyle)
    p.append(pPr)
    
    if bold_prefix:
        # Run en negrita para el prefijo
        r_bold = OxmlElement("w:r")
        rPr_bold = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr_bold.append(b)
        r_bold.append(rPr_bold)
        t_bold = OxmlElement("w:t")
        t_bold.text = bold_prefix
        t_bold.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_bold.append(t_bold)
        p.append(r_bold)
        # Run normal para el resto
        r_normal = OxmlElement("w:r")
        t_normal = OxmlElement("w:t")
        t_normal.text = text
        t_normal.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_normal.append(t_normal)
        p.append(r_normal)
    else:
        r = OxmlElement("w:r")
        t_elem = OxmlElement("w:t")
        t_elem.text = text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t_elem)
        p.append(r)
    
    anchor_elem.addnext(p)
    return p

# Construir los parrafos en orden inverso (cada uno se inserta justo despues del ancla)
lines = [
    # (texto, estilo, bold_prefix)
    # Introduccion
    ("El sistema utiliza MongoDB Atlas como base de datos NoSQL orientada a documentos, organizada en tres colecciones principales que modelan los datos del sistema de gestión de tareas de la Escuela de Posgrado.", "BodyText", None),
    ("", "BodyText", None),
    # --- Coleccion users ---
    ("Coleccion users", "Heading4", None),
    ("Almacena los datos de autenticacion y perfil de cada usuario registrado en el sistema.", "BodyText", None),
    ("\u2022  _id: ObjectId  identificador unico generado por MongoDB.", "ListParagraph", None),
    ("\u2022  nombre: String  nombre completo del usuario.", "ListParagraph", None),
    ("\u2022  email: String (unico)  correo electronico utilizado como credencial de acceso.", "ListParagraph", None),
    ("\u2022  password: String  contrasena encriptada con bcryptjs (10 salt rounds).", "ListParagraph", None),
    ("\u2022  rol: String  define el nivel de acceso: 'Jefa', 'Asistente' o 'Practicante'.", "ListParagraph", None),
    ("\u2022  activo: Boolean  indica si la cuenta esta habilitada.", "ListParagraph", None),
    ("\u2022  createdAt / updatedAt: Date  timestamps automaticos de Mongoose.", "ListParagraph", None),
    ("", "BodyText", None),
    # --- Coleccion tasks ---
    ("Coleccion tasks", "Heading4", None),
    ("Registra todas las tareas asignadas dentro del sistema, con su estado, prioridad y responsable.", "BodyText", None),
    ("\u2022  _id: ObjectId  identificador unico de la tarea.", "ListParagraph", None),
    ("\u2022  titulo: String  nombre descriptivo de la tarea.", "ListParagraph", None),
    ("\u2022  descripcion: String  detalle ampliado del trabajo a realizar.", "ListParagraph", None),
    ("\u2022  estado: String  ciclo de vida de la tarea: 'Pendiente', 'En Proceso' o 'Completada'.", "ListParagraph", None),
    ("\u2022  prioridad: String  nivel de urgencia: 'Alta', 'Media' o 'Baja'.", "ListParagraph", None),
    ("\u2022  responsable: ObjectId  referencia al usuario asignado (coleccion users).", "ListParagraph", None),
    ("\u2022  fechaLimite: Date  fecha maxima de entrega.", "ListParagraph", None),
    ("\u2022  creadoPor: ObjectId  referencia al usuario que creo la tarea (coleccion users).", "ListParagraph", None),
    ("\u2022  createdAt / updatedAt: Date  timestamps automaticos de Mongoose.", "ListParagraph", None),
    ("", "BodyText", None),
    # --- Coleccion notifications ---
    ("Coleccion notifications", "Heading4", None),
    ("Gestiona las notificaciones internas enviadas a los usuarios del sistema.", "BodyText", None),
    ("\u2022  _id: ObjectId  identificador unico de la notificacion.", "ListParagraph", None),
    ("\u2022  destinatario: ObjectId  referencia al usuario receptor (coleccion users).", "ListParagraph", None),
    ("\u2022  tipo: String  categoria de la notificacion: 'tarea_asignada', 'tarea_completada' o 'sistema'.", "ListParagraph", None),
    ("\u2022  mensaje: String  texto descriptivo de la notificacion.", "ListParagraph", None),
    ("\u2022  leida: Boolean  indica si el usuario ya reviso la notificacion (por defecto false).", "ListParagraph", None),
    ("\u2022  fechaCreacion: Date  marca de tiempo de creacion de la notificacion.", "ListParagraph", None),
    ("", "BodyText", None),
    # Relaciones
    ("Relaciones entre colecciones", "Heading4", None),
    ("Las tres colecciones se relacionan mediante referencias de tipo ObjectId: cada tarea referencia a dos usuarios (responsable y creadoPor) de la coleccion users mediante una relacion 1N; cada notificacion referencia a un usuario destinatario de la coleccion users en una relacion 1N. Mongoose gestiona estas referencias mediante populate() para resolver los documentos relacionados en consultas.", "BodyText", None),
    ("", "BodyText", None),
]

# Insertar en orden inverso para mantener la secuencia correcta
last_anchor = anchor_p
for (text, style, bold_prefix) in reversed(lines):
    new_p = add_para_after(last_anchor, text, style, bold_prefix)

print(f"\nInsertados {len(lines)} parrafos del esquema real.")
print(f"Total parrafos: {len(doc.paragraphs)}")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

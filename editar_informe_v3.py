"""
Script v3 — Sección 4.8, 4.9 y secciones finales
  Cambios en:
  · 4.8  Recursos Humanos: tabla software (quitar Multer/Cloudinary/Docker/
          Nodemailer/Nginx/OpenAPI, agregar jsPDF/Recharts/DnD/date-fns)
  · 4.8  Texto intro + Experiencia del equipo
  · 4.9.1 Factibilidad Técnica
  · 4.9.2 Factibilidad Económica (costos: 13 semanas, sin VPS, sin Cloudinary)
  · 4.9.3 Factibilidad Operativa
  · 4.9.4 Factibilidad Legal
  · 4.9.5 Factibilidad Social
  · Beneficios esperados
  · Sección Desarrollo (metodología XP adaptada)
  · Limpieza global residual
"""

from docx import Document

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)

# ─── Utilidades ────────────────────────────────────────────────────────────────
def set_para(para, text):
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ''
    if not para.runs:
        para.add_run(text)

def replace_one(fragment, new_text, min_idx=0):
    key = fragment.lower()[:55]
    for i, para in enumerate(doc.paragraphs):
        if i < min_idx:
            continue
        if key in para.text.lower():
            set_para(para, new_text)
            print(f"    ✓ [{i}] '{fragment[:60]}'")
            return True
    print(f"    ✗ NO ENCONTRADO: '{fragment[:60]}'")
    return False

def replace_all(fragment, new_text):
    key = fragment.lower()[:55]
    count = 0
    for para in doc.paragraphs:
        if key in para.text.lower():
            for run in para.runs:
                if fragment in run.text:
                    run.text = run.text.replace(fragment, new_text)
            count += 1
    if count:
        print(f"    ✓ '{fragment[:55]}' → reemplazado {count}x")
    return count

def replace_in_tables(old, new):
    """Reemplaza texto dentro de celdas de tablas."""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old.lower()[:30] in para.text.lower():
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1
                        # Si no tiene runs, cambiar directamente
                        if old.lower() in para.text.lower() and not para.runs:
                            para.add_run(para.text.replace(old, new))
    return count

# ══════════════════════════════════════════════════════════════════════════════
# 1. TABLA 16 — Requerimientos de Hardware (texto intro)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.8 HARDWARE ---")

replace_one(
    "Se establecen los requerimientos mínimos de hardware nece- sarios para el desarrollo, implementación y funcionamiento",
    "Se establecen los requerimientos mínimos de hardware necesarios para el "
    "desarrollo, implementación y funcionamiento del Sistema de Gestión de "
    "Tareas – Oficina EPO:"
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TABLA 17 — Requerimientos de SOFTWARE
#    Modificar celdas de la tabla en el documento Word
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.8 TABLA SOFTWARE ---")

replace_one(
    "Se determinan los requerimientos de software necesarios para el desarrollo e implementación del sistema:",
    "Se determinan los requerimientos de software necesarios para el desarrollo "
    "e implementación del Sistema de Gestión de Tareas – Oficina EPO:"
)

# Reemplazos dentro de las tablas
tabla_software = [
    # (viejo, nuevo)  — se busca en celdas
    ("Formularios", "Generación PDF"),
    ("RHF + Zod", "jsPDF + jsPDF-autotable"),
    ("React Hook Form y Zod para validación de formularios.", 
     "Generación automática de reportes PDF tabulares con jsPDF y jsPDF-autotable."),
    ("React Hook Form y Zod para validación de for-\nmularios.",
     "Generación automática de reportes PDF tabulares con jsPDF y jsPDF-autotable."),
    ("Multer", "Recharts"),
    ("Manejo de uploads en el backend.",
     "Biblioteca de gráficos para React (barras, líneas) para el dashboard."),
    ("Cloudinary", "React Beautiful DnD"),
    ("Almacenamiento/CDN de imágenes y recursos.",
     "Arrastrar y soltar accesible para reorganizar tareas en la interfaz."),
    ("Correo", "Fechas"),
    ("Nodemailer", "date-fns + Moment.js"),
    ("Envío de emails vía SMTP.", "Manejo, formateo y manipulación de fechas en frontend."),
    ("Envío de emails vía SMTP. Config/seguridad",
     "Manejo, formateo y manipulación de fechas en frontend."),
    ("Contenedores", "Despliegue"),
    ("Docker + Compo-\nse", "Render + Vercel"),
    ("Docker + Compose", "Render + Vercel"),
    ("Empaquetado y despliegue reproducible.", 
     "Plataformas cloud gratuitas: backend en Render, frontend en Vercel."),
    ("Servidor web", "Repositorio"),
    ("Nginx (proxy)", "Git + GitHub"),
    ("Proxy reverso y estáticos en producción.",
     "Control de versiones y colaboración del código fuente."),
    ("Documentación API", "Iconos UI"),
    ("OpenAPI 3 (opcio-\nnal)", "Lucide React"),
    ("OpenAPI 3 (opcional)", "Lucide React"),
    ("Especificación de endpoints y contratos de la API.",
     "Biblioteca de iconos SVG para la interfaz de usuario React."),
    ("Subida de archi-\nvos", "Drag & Drop"),
    ("Subida de archivos", "Drag & Drop"),
]

tc = 0
for old, new in tabla_software:
    n = replace_in_tables(old, new)
    if n:
        print(f"    ✓ Tabla: '{old[:45]}' → '{new[:35]}' ({n}x)")
        tc += n
print(f"  Total celdas modificadas: {tc}")

# ══════════════════════════════════════════════════════════════════════════════
# 3. EXPERIENCIA DEL EQUIPO
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- EXPERIENCIA EQUIPO ---")

replace_one(
    "El practicante preprofesional asignado al proyecto posee forma- ción académica y experiencia práctica en múltiples áreas",
    "El practicante preprofesional asignado al proyecto posee formación académica "
    "y experiencia práctica en múltiples áreas tecnológicas relevantes, lo que "
    "asegura un enfoque técnico sólido para la ejecución exitosa del proyecto. "
    "El conocimiento abarca desarrollo web full-stack con JavaScript, arquitectura "
    "de aplicaciones cliente-servidor con stack MERN (MongoDB, Express, React, "
    "Node.js), bases de datos NoSQL, generación de documentos PDF con jsPDF, "
    "visualización de datos con Recharts, metodologías ágiles de desarrollo (XP), "
    "control de versiones con Git y despliegue en plataformas cloud (Render y "
    "Vercel), garantizando capacidad para implementar todas las fases del ciclo "
    "de vida del sistema propuesto."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. FACTIBILIDAD TÉCNICA (4.9.1)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.9.1 FACTIBILIDAD TÉCNICA ---")

replace_one(
    "Sistema Operativo:",
    "Sistema Operativo:"
)

replace_one(
    "Sistema de Gestión de Tareas – Oficina EPO, una aplicación web full-stack basada en arquitectura cliente-servidor con st",
    "Con el fin de asegurar la compatibilidad y el rendimiento eficiente del "
    "Sistema de Gestión de Tareas – Oficina EPO, se ha considerado la "
    "compatibilidad con Windows 10/11 para estaciones de trabajo administrativas "
    "y el uso de plataformas cloud (Render para el backend y Vercel para el "
    "frontend) que no requieren configuración de servidor dedicado por parte "
    "del equipo de prácticas.",
    min_idx=4000
)

replace_one(
    "El lenguaje de programación seleccionado debe cumplir con cri- terios como la capacidad para manejar múltiples conexione",
    "El lenguaje de programación seleccionado debe cumplir con criterios como la "
    "capacidad para manejar múltiples operaciones asíncronas, facilidad en el "
    "desarrollo y mantenimiento, así como estabilidad y robustez. Por ello, se "
    "ha decidido utilizar JavaScript con Node.js y Express para el backend, y "
    "React 18 con Vite para el frontend. Esta combinación tecnológica permite "
    "desarrollo ágil, ciclos de entrega cortos, generación de reportes PDF con "
    "jsPDF y facilita la integración mediante API REST.",
    min_idx=4000
)

replace_one(
    "La elección de estas tecnologías responde a su capacidad para operar tanto de forma integrada con sistemas existentes en",
    "La elección de estas tecnologías responde a su capacidad para operar de "
    "forma independiente mediante navegador web estándar. El sistema ha sido "
    "diseñado con arquitectura modular que facilita su mantenimiento evolutivo "
    "y permite incorporar futuras funcionalidades sin rediseños estructurales "
    "significativos.",
    min_idx=4000
)

replace_one(
    "La selección del sistema gestor de bases de datos es funda- mental para garantizar la flexibilidad de esquema, escalabil",
    "La selección del sistema gestor de bases de datos es fundamental para "
    "garantizar la flexibilidad de esquema, escalabilidad y rendimiento del "
    "sistema. Se utiliza MongoDB en su versión 6.0 o superior, aprovechando "
    "el servicio gestionado MongoDB Atlas (plan gratuito) que ofrece respaldos "
    "automatizados, réplicas y escalabilidad sin intervención manual. El modelado "
    "de datos se realiza mediante Mongoose, lo que proporciona validaciones de "
    "esquema y facilita el mantenimiento del código de acceso a datos. Las tres "
    "colecciones del sistema son: usuarios, tareas y notificaciones.",
    min_idx=4000
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. FACTIBILIDAD ECONÓMICA (4.9.2) — Costos actualizados
#    13 semanas (15/09 al 14/12), 4h diarias, S/.8.50/h
#    260h * S/.8.50 = S/.2,210
#    Sin Cloudinary, sin VPS, todo gratuito
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.9.2 FACTIBILIDAD ECONÓMICA ---")

replace_one(
    "El desarrollo del sistema es ejecutado por un único practican- te preprofesional de la carrera de Ingeniería de Sistemas",
    "El desarrollo del sistema es ejecutado por un único practicante "
    "preprofesional de la carrera de Ingeniería de Sistemas como parte de sus "
    "prácticas preprofesionales obligatorias (período 15/09/2025 – 14/12/2025). "
    "La propuesta no contempla contrataciones adicionales de personal externo; "
    "la supervisión y validación institucional es proporcionada por la "
    "Coordinadora del CMC-EPO sin generar costo directo para el proyecto. "
    "Para efectos de evaluación económica se considera una asignación económica "
    "referencial equivalente al valor de mercado de un practicante en desarrollo "
    "de software.",
    min_idx=4000
)

# Nota del cálculo de costos de personal
replace_one(
    "Nota. Cálculo referencial: 16 semanas de duración del proyecto, 4 horas diarias promedio dedicadas al desarrollo y prueb",
    "Nota. Cálculo referencial: 13 semanas de duración del proyecto "
    "(15/09/2025 – 14/12/2025), 4 horas diarias promedio dedicadas al "
    "desarrollo y pruebas, costo por hora de practicante S/. 8.50 según mercado "
    "local. Elaboración propia.",
    min_idx=4000
)

# Actualizar el total en la tabla de costos: 16sem*4h*5d*S/.8.50 = S/.2,720
# → 13sem*4h*5d*S/.8.50 = S/.2,210
# Buscar en tablas
n = replace_in_tables("2,720.00", "2,210.00")
print(f"    ✓ Monto 2,720 → 2,210: {n}x")
n = replace_in_tables("2,720", "2,210")
print(f"    ✓ Monto 2720 → 2210: {n}x")
n = replace_in_tables("S/. 680", "S/. 170")   # mensual = 4h*5d*4semanas*8.50 ≈ 680 → para 13 sem es referencial
# En realidad voy a dejarlo como mensual S/.680 pero cambiar total a 2210 y semanas a 13
n = replace_in_tables("16 semanas", "13 semanas")
print(f"    ✓ 16 semanas → 13 semanas: {n}x")

replace_one(
    "Total costos de personal 2,720.00",
    "Total costos de personal              S/. 2,210.00",
    min_idx=4000
)

# Tabla 22 - Costo de desarrollo: quitar VPS 240, quitar Cloudinary
replace_one(
    "El desarrollo del sistema se realizó aprovechando tecnologías open source y servicios en la nube con planes iniciales gr",
    "El desarrollo del sistema se realizó aprovechando íntegramente tecnologías "
    "open source y servicios en la nube con planes gratuitos. El frontend se "
    "desplegó en Vercel (plan gratuito), el backend en Render (plan gratuito), "
    "la base de datos en MongoDB Atlas (plan gratuito M0) y el dominio "
    "provisional utiliza el subdominio provisto por dichas plataformas. "
    "No se incurrió en costos de infraestructura durante el período de prácticas.",
    min_idx=4000
)

# Actualizar total en Tabla 22
n = replace_in_tables("Servidor VPS básico (4 meses)", "Despliegue backend Render (plan gratuito)")
n += replace_in_tables("240.00", "0.00")
n += replace_in_tables("Dominio (.pe o .com) anual", "Despliegue frontend Vercel (plan gratuito)")
n += replace_in_tables("60.00", "0.00")
n += replace_in_tables("CDN y almacenamiento de medios Cloudinary (plan", "Iconos y bibliotecas UI (open source)")
print(f"    ✓ Tabla costos: {n} cambio(s)")

replace_one(
    "Total 3,020.00",
    "Total              S/. 2,210.00",
    min_idx=4000
)
# También en tabla
n = replace_in_tables("3,020.00", "2,210.00")
n += replace_in_tables("3,020", "2,210")
print(f"    ✓ Total 3,020 → 2,210: {n}x")

replace_one(
    "Nota. Montos referenciales orientados a un proyecto académico de prácticas preprofesionales aprovechando infraestructura",
    "Nota. Montos referenciales orientados a un proyecto académico de prácticas "
    "preprofesionales aprovechando íntegramente infraestructura cloud gratuita "
    "y tecnologías open source. Elaboración propia.",
    min_idx=4000
)

# Costo de software párrafo
replace_one(
    "Dado que el sistema utiliza exclusivamente tecnologías open source (Node.js, Express, React, Vite, Tailwind CSS, Mongoos",
    "Dado que el sistema utiliza exclusivamente tecnologías open source (Node.js, "
    "Express, React, Vite, Tailwind CSS, Mongoose, MongoDB, jsPDF, Recharts, "
    "React Beautiful DnD, date-fns, Moment.js) y servicios en la nube con planes "
    "gratuitos suficientes para el volumen operativo proyectado, los gastos "
    "asociados al licenciamiento de software son nulos. Todas las herramientas "
    "de desarrollo empleadas (Visual Studio Code, Git, GitHub) son de uso libre "
    "sin restricciones comerciales.",
    min_idx=4000
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. FACTIBILIDAD OPERATIVA (4.9.3)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.9.3 FACTIBILIDAD OPERATIVA ---")

replace_one(
    "La viabilidad operativa del Sistema Integral de Gestión Acadé- mica (SIGA) propuesto para la Escuela Profesional de Odontología",
    "La viabilidad operativa del Sistema de Gestión de Tareas – Oficina EPO "
    "para el Comité de Mejora Continua de la Escuela Profesional de Odontología "
    "de la Universidad Privada de Tacna se fundamenta en varios aspectos críticos "
    "que aseguran su efectiva implementación y operación sostenible dentro del área.",
    min_idx=4100
)

replace_one(
    "El sistema formaliza y centraliza procesos que actualmente se ejecutan de manera manual y dispersa: difusión de noticias",
    "El sistema formaliza y centraliza procesos que actualmente se ejecutan de "
    "manera manual y dispersa: registro y seguimiento de tareas administrativas, "
    "control de estado de actividades (pendiente/completada), generación de "
    "reportes mensuales de productividad y gestión de notificaciones internas. "
    "La consolidación de estos flujos en un panel único reduce significativamente "
    "los pasos manuales, elimina la dispersión de información y facilita la "
    "trazabilidad completa de las actividades del área.",
    min_idx=4100
)

replace_one(
    "La escuela profesional cuenta con el personal administrativo adecuadoparaoperarelsistemasinnecesidadderealizarcontratacio",
    "El área cuenta con el personal adecuado para operar el sistema sin "
    "necesidad de contrataciones adicionales. Se asigna el rol de Jefa a la "
    "coordinadora del CMC con acceso completo a todas las funcionalidades "
    "(CRUD de tareas, reportes, estadísticas, notificaciones, gestión de "
    "usuarios), mientras que el rol de Asistente permite el registro y completado "
    "de tareas asignadas. La carga operativa se simplifica mediante filtros por "
    "estado y prioridad, búsquedas textuales y acceso desde cualquier dispositivo "
    "con navegador web.",
    min_idx=4100
)

replace_one(
    "Se espera que el sistema aumente significativamente la produc- tividad del personal administrativo",
    "Se espera que el sistema aumente significativamente la productividad del "
    "personal al optimizar procesos clave. El registro de tareas, que "
    "anteriormente requería anotaciones físicas y hojas de cálculo sin "
    "integración, se convierte en una operación de segundos ejecutable desde "
    "cualquier dispositivo. La generación de reportes mensuales en PDF, que "
    "anteriormente demandaba horas de consolidación manual, se reduce a un "
    "proceso automático de selección de período y descarga inmediata. El "
    "dashboard con estadísticas en tiempo real elimina la necesidad de cálculos "
    "manuales para evaluar la productividad del área.",
    min_idx=4100
)

replace_one(
    "Para garantizar la adopción efectiva del sistema se contempla una inducciónoperativabreve",
    "Para garantizar la adopción efectiva del sistema se contempla una inducción "
    "operativa breve (2 a 3 horas) dirigida a la Jefa y Asistente del CMC, "
    "acompañada de manual de usuario conciso con capturas de pantalla y "
    "explicaciones paso a paso. La interfaz web es intuitiva, responsive y sigue "
    "patrones de usabilidad establecidos, lo que facilita la curva de aprendizaje. "
    "No se anticipan impactos negativos sobre otros sistemas institucionales dado "
    "que el sistema opera de forma independiente mediante navegador web estándar.",
    min_idx=4100
)

replace_one(
    "Las tareas operativas periódicas incluyen revisión de respaldos automatizadosdebasededatos",
    "Las tareas operativas periódicas incluyen revisión de respaldos automáticos "
    "de MongoDB Atlas, actualización de dependencias en ventanas de mantenimiento "
    "planificadas, monitoreo de disponibilidad del servicio en Render y Vercel, "
    "y verificación semestral de cuentas de usuario y permisos por rol.",
    min_idx=4100
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FACTIBILIDAD LEGAL (4.9.4)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.9.4 FACTIBILIDAD LEGAL ---")

replace_one(
    "El análisis de viabilidad legal del sistema se enfoca en asegurar el cumplimiento riguroso con las normativas peruanas",
    "El análisis de viabilidad legal del Sistema de Gestión de Tareas – Oficina EPO "
    "se enfoca en asegurar el cumplimiento riguroso con las normativas peruanas "
    "vigentes en materia de protección de datos personales, propiedad intelectual "
    "y transparencia institucional, garantizando que la operación del sistema sea "
    "legal, segura y éticamente responsable.",
    min_idx=4100
)

replace_one(
    "El sistema cumple estrictamente con la Ley N.º 29733 - Ley de Protección de Datos Personales del Perú",
    "El sistema cumple con la Ley N.° 29733 – Ley de Protección de Datos "
    "Personales del Perú: recolecta únicamente los datos estrictamente necesarios "
    "(nombre, usuario, contraseña cifrada) para la autenticación y gestión de "
    "tareas, sin solicitar información sensible o irrelevante. Los usuarios pueden "
    "solicitar rectificación o eliminación de sus datos mediante procedimiento "
    "documentado. Las contraseñas se almacenan exclusivamente como hash bcrypt "
    "(cost factor 10), garantizando que no puedan recuperarse en texto plano.",
    min_idx=4100
)

replace_one(
    "La operación del sistema se alinea con principios de privacidad desde el diseño",
    "La operación del sistema se alinea con principios de privacidad desde el "
    "diseño: cifrado de datos en tránsito mediante HTTPS con certificados TLS "
    "automáticos provistos por Render y Vercel, control de acceso basado en "
    "roles con autenticación JWT, y almacenamiento seguro de credenciales "
    "mediante hashing bcrypt. El código fuente desarrollado es propiedad "
    "intelectual de la institución académica conforme al marco normativo de "
    "prácticas preprofesionales.",
    min_idx=4100
)

replace_one(
    "Cuando el sistema hace uso de servicios de terceros para pro- cesamiento o almacenamiento de datos (MongoDB Atlas",
    "El sistema utiliza exclusivamente tecnologías open source con licencias "
    "permisivas (MIT, BSD) y servicios cloud gratuitos (MongoDB Atlas, Render, "
    "Vercel) con términos de servicio compatibles con uso académico. Todo el "
    "código fuente desarrollado es propiedad de la institución y puede ser "
    "auditado, modificado y desplegado libremente por el personal técnico "
    "autorizado.",
    min_idx=4100
)

replace_one(
    "Todo el desarrollo del sistema utiliza exclusivamente tecnolo- gías open source con licencias compatibles y permisivas",
    "Las tecnologías utilizadas cuentan con licencias open source compatibles: "
    "Node.js (MIT), Express (MIT), React (MIT), Vite (MIT), Tailwind CSS (MIT), "
    "MongoDB Community (SSPL), jsPDF (MIT), Recharts (MIT), React Beautiful DnD "
    "(Apache 2.0). No se incurre en infracciones de derechos de autor ni en uso "
    "no autorizado de software propietario.",
    min_idx=4100
)

replace_one(
    "El sistema provee términos de uso claros y política de priva- cidad exhaustiva",
    "El sistema implementa buenas prácticas de seguridad: tokens JWT con "
    "expiración configurable, validación de entradas en frontend y backend, "
    "manejo centralizado de errores sin exposición de información sensible en "
    "respuestas de la API, y variables de entorno almacenadas en archivo .env "
    "excluido del control de versiones (.gitignore).",
    min_idx=4100
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. FACTIBILIDAD SOCIAL (4.9.5)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.9.5 FACTIBILIDAD SOCIAL ---")

replace_one(
    "La factibilidad social del Sistema Integral de Gestión Acadé- mica (SIGA) para la Escuela Profesional",
    "La factibilidad social del Sistema de Gestión de Tareas – Oficina EPO "
    "es alta, dado que el proyecto responde directamente a necesidades reales "
    "identificadas durante el período de prácticas preprofesionales en el área "
    "de gestión administrativa del Comité de Mejora Continua. El sistema aporta "
    "beneficios tangibles para el personal del área (Jefa y Asistente del CMC).",
    min_idx=4100
)

replace_one(
    "Para los usuarios internos (personal administrativo), el sistema reducesignificativamentetiemposdedicadosatareasrepetitivas",
    "Para los usuarios del sistema (Jefa y Asistente del CMC), el sistema "
    "reduce significativamente el tiempo dedicado a tareas repetitivas como el "
    "registro manual de actividades en hojas de cálculo, la elaboración de "
    "informes mensuales y el seguimiento informal del estado de trabajos "
    "pendientes. La centralización de tareas en un panel único mejora la "
    "transparencia y la rendición de cuentas sobre actividades realizadas, "
    "facilitando la generación de reportes PDF para evaluación de la gestión "
    "del área y toma de decisiones basada en datos verificables.",
    min_idx=4100
)

replace_one(
    "Para los usuarios externos (comunidad universitaria y público general)",
    "El diseño del sistema fue validado con la coordinadora del área mediante "
    "sesiones de levantamiento de requisitos y retroalimentación continua durante "
    "el desarrollo iterativo, lo que favorece la adopción al garantizar que las "
    "funcionalidades implementadas respondan a necesidades reales expresadas por "
    "los usuarios finales. La interfaz web es responsive, accesible desde "
    "escritorio, tablet y móvil, y sigue patrones de usabilidad que facilitan "
    "la curva de aprendizaje y reducen la resistencia al cambio tecnológico.",
    min_idx=4100
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. BENEFICIOS ESPERADOS (sección dentro de 4.6.4)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- BENEFICIOS ESPERADOS ---")

replace_one(
    "Eficiencia Operativa: Automatización de procesos administra- tivos que libera 15-20 horas semanales del person",
    "Eficiencia Operativa: Automatización del registro de tareas y generación "
    "de reportes que libera 3-5 horas semanales del personal, redirigiendo "
    "esfuerzos hacia actividades de mayor valor para el CMC."
)

replace_one(
    "Centralización de Información: Repositorio digital único acce- sible 24/7 desde cualquier dispositivo. Elimina",
    "Centralización de Información: Repositorio digital único accesible 24/7 "
    "desde cualquier dispositivo. Elimina registros dispersos en hojas de "
    "cálculo y anotaciones físicas, garantizando trazabilidad completa."
)

replace_one(
    "Mejora en Toma de Decisiones: Reportes estadísticos auto- máticos sobre captación, participación y admisiones.",
    "Mejora en Toma de Decisiones: Reportes PDF mensuales automáticos y "
    "dashboard con estadísticas en tiempo real que eliminan la consolidación "
    "manual de datos, reduciendo el tiempo de elaboración de informes de "
    "horas a segundos."
)

replace_one(
    "Calidad de Servicio: Mejora en experiencia estudiantil median- te acceso oportuno a información y eventos.",
    "Calidad de Gestión: Mejora en la experiencia del personal mediante "
    "acceso inmediato al estado de todas las tareas, notificaciones internas "
    "y visualización de productividad histórica del área."
)

replace_one(
    "Escalabilidad: Arquitectura modular API RESTful para agregar funcionalidades. MongoDB flexible para evolución",
    "Escalabilidad: Arquitectura modular con API RESTful que permite agregar "
    "futuras funcionalidades (asignación masiva, calendario, integración con "
    "otros sistemas) sin rediseños estructurales. MongoDB flexible para "
    "evolución del modelo de datos."
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. SECCIÓN DESARROLLO — metodología XP (párrafo de apertura)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- DESARROLLO / METODOLOGÍA ---")

replace_one(
    "Para el desarrollo del SIGA la Oficina EPO se adoptó Extreme Program- ming (XP) como metodología ágil. XP prio",
    "Para el desarrollo del Sistema de Gestión de Tareas – Oficina EPO se "
    "adoptó Extreme Programming (XP) como metodología ágil. XP prioriza la "
    "entrega temprana y frecuente de software funcional, la simplicidad del "
    "diseño y la adaptación continua al cambio, validando primero los flujos "
    "de mayor valor (gestión de tareas, autenticación por roles, generación "
    "de reportes PDF) y ajustando reglas de negocio conforme surgían nuevos "
    "requerimientos durante el período de prácticas."
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. LIMPIEZA GLOBAL RESIDUAL en párrafos
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- LIMPIEZA RESIDUAL GLOBAL ---")

residuales_paras = [
    ("SIGA Odontología",                    "Sistema de Gestión de Tareas – Oficina EPO"),
    ("SIGA la Oficina EPO",                 "Sistema de Gestión de Tareas – Oficina EPO"),
    ("Sistema Integral de Gestión Académica","Sistema de Gestión de Tareas"),
    ("gestión académica",                   "gestión de tareas"),
    ("prospectos académicos",               "tareas administrativas"),
    ("noticias institucionales",            "reportes mensuales"),
    ("gestión de eventos",                  "gestión de tareas"),
    ("inscripciones",                       "registro de tareas"),
    ("estudiantes potenciales",             "personal del área"),
    ("administrador, docente, estudiante",  "Jefa y Asistente"),
    ("Escuela Profesional de Odontología",  "Oficina EPO – CMC"),
    ("Docker",                              "Render/Vercel"),
    ("Cloudinary",                          "MongoDB Atlas"),
    ("Multer",                              "jsPDF"),
    ("Nodemailer",                          "date-fns"),
    ("leads",                               "tareas"),
]

for old, new in residuales_paras:
    for para in doc.paragraphs:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

# También en tablas
residuales_tablas = [
    ("SIGA Odontología",    "Sistema Gestión Tareas EPO"),
    ("Multer",              "jsPDF"),
    ("Cloudinary",          "Recharts"),
    ("Nodemailer",          "date-fns"),
    ("Docker + Compose",    "Render + Vercel"),
    ("Docker + Compo-",     "Render +"),
    ("Nginx (proxy)",       "Git + GitHub"),
    ("OpenAPI 3",           "Lucide React"),
    ("RHF + Zod",           "jsPDF-autotable"),
]
for old, new in residuales_tablas:
    n = replace_in_tables(old, new)
    if n:
        print(f"    ✓ Tabla residual: '{old}' → '{new}' ({n}x)")

print("    ✓ Limpieza completada")

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
doc.save(DEST)
print("\n✅ INFORME.docx v3 guardado con todos los cambios de 4.8 y 4.9.")

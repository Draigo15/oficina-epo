import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from lxml import etree
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)
p    = doc.paragraphs

def ft(para):
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))

def set_xml(para, new_text):
    all_t = para._p.findall(f'.//{{{NS}}}t')
    if all_t:
        all_t[0].text = new_text
        for t in all_t[1:]: t.text = ''
    else:
        para.add_run(new_text)

def rep_xml(para, old, new):
    full = ft(para)
    if old in full:
        set_xml(para, full.replace(old, new, 1))
        return True
    return False

# ── [39]  Índice TOC: Caso de Uso CS-02 Noticias ────────────────────────────
rep_xml(p[39],
    'Narrativa del Caso de Uso CS-02 – Consultar Noticias Institucionales',
    'Narrativa del Caso de Uso CS-02 – Consultar Tareas Asignadas')
print(f'[39] → {ft(p[39])[:80]}')

# ── [44]  Índice TOC: Caso de Uso CS-07 Prospectos ──────────────────────────
rep_xml(p[44],
    'Narrativa del Caso de Uso CS-07 – Gestionar Prospectos Académicos',
    'Narrativa del Caso de Uso CS-07 – Gestionar Notificaciones')
print(f'[44] → {ft(p[44])[:80]}')

# ── [45]  Índice TOC: Caso de Uso CS-08 Noticias ────────────────────────────
rep_xml(p[45],
    'Narrativa del Caso de Uso CS-08 – Gestionar Noticias Institucionales',
    'Narrativa del Caso de Uso CS-08 – Generar Reportes PDF')
print(f'[45] → {ft(p[45])[:80]}')

# ── [130] Índice: Panel de gestión de noticias ───────────────────────────────
rep_xml(p[130],
    'Panel de gestión de noticias - vista administrativa',
    'Dashboard de estadísticas – gráficos Recharts')
print(f'[130] → {ft(p[130])[:80]}')

# ── [371] Párrafo que menciona "eventos académicos" ──────────────────────────
rep_xml(p[371],
    'eventos académicos',
    'actividades del área')
print(f'[371] → {ft(p[371])[:80]}')

# ── [841]  Narrativa CU CS-02 (en el cuerpo) ────────────────────────────────
rep_xml(p[841],
    'Narrativa del Caso de Uso CS-02 – Consultar Noticias Institucionales',
    'Narrativa del Caso de Uso CS-02 – Consultar Tareas Asignadas')
print(f'[841] → {ft(p[841])[:80]}')

# ── [1562] Narrativa CU CS-07 (en el cuerpo) ────────────────────────────────
rep_xml(p[1562],
    'Narrativa del Caso de Uso CS-07 – Gestionar Prospectos Académicos',
    'Narrativa del Caso de Uso CS-07 – Gestionar Notificaciones')
print(f'[1562] → {ft(p[1562])[:80]}')

# ── [1831] Narrativa CU CS-08 (en el cuerpo) ────────────────────────────────
rep_xml(p[1831],
    'Narrativa del Caso de Uso CS-08 – Gestionar Noticias Institucionales',
    'Narrativa del Caso de Uso CS-08 – Generar Reportes PDF')
print(f'[1831] → {ft(p[1831])[:80]}')

# ── [1852] "Gestión de Noticias" en texto de caso de uso ────────────────────
rep_xml(p[1852],
    'Gestión de Noticias',
    'Gestión de Reportes')
print(f'[1852] → {ft(p[1852])[:80]}')

# ── [3757] Diagrama de despliegue — Cloudinary / CDN ────────────────────────
# Es texto concatenado: la referencia a "CDN / Static Hosting" con "Cloudinary"
rep_xml(p[3757],
    'CDN / Static Hosting',
    'Vercel / Static Hosting')
rep_xml(p[3757],
    'Cloudinary',
    'MongoDB Atlas')
print(f'[3757] → {ft(p[3757])[:100]}')

# ── Corregir también los 3 headers de Anexos que quedaron sin procesar ───────
# [4511] "Perfil de Administrador", [4526] "Módulo del Estudiante", [4557] "Módulo del Docente"
for idx, old, new in [
    (4511, 'Perfil de Administrador',  'Perfil de Usuario – Administrador'),
    (4526, 'Módulo del Estudiante',    'Módulo de Notificaciones'),
    (4557, 'Módulo del Docente',       'Módulo de Reportes PDF'),
]:
    # Buscar desde idx-2 hasta idx+5
    for i in range(max(0, idx-5), min(len(p), idx+10)):
        if old in ft(p[i]):
            set_xml(p[i], new)
            print(f'[{i}] header anexo: "{old}" → "{new}"')
            break

# ── Corregir [4302] y [4308] conclusiones cortadas ────────────────────────────
for idx in [4302, 4308]:
    txt = ft(p[idx])
    if txt.strip():
        set_xml(p[idx], '')
        print(f'[{idx}] conclusión cortada vaciada: {txt[:60]}')

doc.save(DEST)
print('\n✅ Correcciones finales aplicadas y guardadas.')

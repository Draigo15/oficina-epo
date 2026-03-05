# -*- coding: utf-8 -*-
# Paso 1: Eliminar bloque de narrativas del portal [649..2453]
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

# Verificar limites antes de eliminar
print(f"Total parrafos: {len(paras)}")
print(f"[649] = '{paras[649].text[:60]}'")
print(f"[2453] = '{paras[2453].text[:60]}'")
print(f"[2454] = '{paras[2454].text[:60]}'")
print()

# Eliminar de atras hacia adelante [649..2453] inclusive
count = 0
for idx in range(2453, 648, -1):  # de 2453 a 649 inclusive
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)
    count += 1

print(f"Eliminados: {count} parrafos")
print(f"Total tras eliminar: {len(doc.paragraphs)}")
print(f"[649] nuevo = '{doc.paragraphs[649].text[:80]}'")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

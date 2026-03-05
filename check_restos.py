# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME_FINAL_SISTEMA.docx")

errores_encontrados = []
palabras_base = ["noticia", "institucional", "estudiante", "contacto", "programa", "prospecto", "autenticar usua-", "conte-"]
palabras_cortadas_regex = re.compile(r'\b[a-zA-ZáéíóúÁÉÍÓÚñÑ]+-\s*\n?\s*[a-zA-ZáéíóúÁÉÍÓÚñÑ]+\b')

for i, p in enumerate(doc.paragraphs):
    text = p.text.lower()
    
    # Buscar palabras base
    for bt in palabras_base:
        if bt in text and len(text.strip()) > 0:
            errores_encontrados.append(f"Párrafo {i} contiene '{bt}': {p.text.strip()[:80]}...")
            
    # Buscar palabras cortadas con guión
    cortes = palabras_cortadas_regex.findall(p.text)
    if cortes:
        errores_encontrados.append(f"Párrafo {i} tiene palabras cortadas: {cortes}")

# Buscar en tablas
for i, table in enumerate(doc.tables):
    for j, row in enumerate(table.rows):
        for k, cell in enumerate(row.cells):
            text = cell.text.lower()
            for bt in palabras_base:
                if bt in text and len(text.strip()) > 0:
                    errores_encontrados.append(f"Tabla {i}, Fila {j} contiene '{bt}': {cell.text.strip()[:80]}...")
            
            cortes = palabras_cortadas_regex.findall(cell.text)
            if cortes:
                errores_encontrados.append(f"Tabla {i}, Fila {j} tiene palabras cortadas: {cortes}")

if not errores_encontrados:
    print("El documento está limpio. No se encontraron restos del sistema base ni palabras cortadas.")
else:
    print(f"Se encontraron {len(errores_encontrados)} posibles errores:")
    # deduplicar e imprimir los primeros 20
    for e in list(set(errores_encontrados))[:20]:
        print("-", e)

# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME.docx")

def clean_paragraph(p):
    # Remueve espacios en blanco dobles dentro de los runs
    for r in p.runs:
        if r.text:
            cleaned = re.sub(r' +', ' ', r.text)
            r.text = cleaned

for p in doc.paragraphs:
    # 1. Año en carátula
    if "Tacna" in p.text and "2025" in p.text:
        for r in p.runs:
            r.text = r.text.replace("2025", "2026")
            
    # 2. Arreglar "Progra-" [salto] "mas" -> Es mejor hacerlo a nivel texto de párrafo, 
    # pero eso rompe runs. Sin embargo, para cosas chicas es aceptable si reconstruimos.
    # Como es riesgoso perder formato, solo hacemos limpieza básica
    clean_paragraph(p)
    
# Reparar tablas que puedan tener textos cortados
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                clean_paragraph(p)

doc.save("INFORME_Limpio.docx")
print("INFORME_Limpio guardado")

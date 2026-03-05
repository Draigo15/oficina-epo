# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME_FINAL_SISTEMA.docx")

def safe_replace(para):
    # This function safely reunites hyphenated words in runs
    text = para.text
    if "-" in text:
        # Check specifically for "funciona- \n lidad" or "funciona-lidad" 
        # which can span across word boundaries with spaces
        fixed = re.sub(r'([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)-\s*\n?\s*([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)', r'\1\2', text)
        if fixed != text:
             # We clear the paragraph and reset the text to unbreak it completely
             # (clearing runs loses some inline formatting but guarantees text continuity)
             para.text = fixed

for p in doc.paragraphs:
    safe_replace(p)

    # 2. Fix the specific break shown in the image: "funciona-"
    if "funciona-" in p.text.lower():
         p.text = p.text.replace("funciona-", "funcionalidades")

# Check tables too because sometimes text gets boxed
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
             for p in cell.paragraphs:
                 safe_replace(p)

doc.save("INFORME_FINAL_SISTEMA_fixed.docx")
print("Arreglado el texto cortado como 'funciona-' y otros.")

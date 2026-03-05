# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

to_delete = []

for i, p in enumerate(paras):
    t = p.text.strip()
    if not t:
        continue

    # 1. Texto espaciado tipo "u s e r s" o "A u te n t ic"
    words = t.split(" ")
    if len(words) >= 4 and all(len(w) <= 3 for w in words) and len(t) > 5:
        to_delete.append((i, "ESPACIADO", t[:60]))
        continue

    # 2. Fragmentos sueltos Heading 3 que son silabas/restos
    if p.style.name == "Heading 3" and t in ["tacto", "Admin"]:
        to_delete.append((i, "FRAGMENTO", t))
        continue

print(f"Parrafos a eliminar: {len(to_delete)}")
for idx, tipo, text in to_delete:
    print(f"  [{idx}] [{tipo}] '{text}'")

# Eliminar de atras hacia adelante para no alterar indices
for idx, tipo, text in sorted(to_delete, key=lambda x: x[0], reverse=True):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)

print(f"\n{len(to_delete)} parrafos eliminados.")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Archivo bloqueado, guardado en {out}")

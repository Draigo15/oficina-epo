# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_Limpio.docx")

keywords_to_remove = ["Noticia", "Programa", "Prospecto", "Lead", "Mensaje", "Contacto", "Estudiante", "Contenido"]

elim_tables = 0
for table in doc.tables:
    text_in_first_row = ""
    try:
        text_in_first_row = table.rows[0].cells[0].text + " " + table.rows[0].cells[1].text
    except:
        pass
        
    for kw in keywords_to_remove:
        if kw.lower() in text_in_first_row.lower():
            table._element.getparent().remove(table._element)
            elim_tables += 1
            break

# Borrar parrafos
elim_p = 0
for p in doc.paragraphs:
    text = p.text.lower()
    if any(kw.lower() in text for kw in keywords_to_remove):
        # Asegurarnos de no borrar palabras comunes como "programa" que puedan ser del lenguaje, 
        # pero para títulos de casos de uso aplican.
        if "caso de uso" in text or "cs-" in text:
            p._element.getparent().remove(p._element)
            elim_p += 1
            

doc.save("INFORME_FINAL_SISTEMA.docx")
print(f"Borradas {elim_tables} tablas y {elim_p} parrafos")

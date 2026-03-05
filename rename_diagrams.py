# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA.docx")

reemplazos = {
    "Diagrama de Secuencia CS-04  Enviar Mensaje de Contacto": "Diagrama de Secuencia CS-04  Gestionar Notificaciones",
    "Diagrama de Secuencia CS-04  Enviar Mensaje de Con-": "Diagrama de Secuencia CS-04  Gestionar Notificaciones",
    "Diagrama de Secuencia CS-08  Gestionar Leads": "Diagrama de Secuencia CS-02  Gestionar Tareas",
    "Diagrama de Secuencia CS-09  Gestionar Contenido": "Diagrama de Secuencia CS-03  Generar Reporte PDF",
    # Índice (cabeceras del índice de figuras)
    "Diagrama de Secuencia CS-04  Enviar Mnesaje de Contacto": "Diagrama de Secuencia CS-04  Gestionar Notificaciones",
}

cambiados = 0
for p in doc.paragraphs:
    for viejo, nuevo in reemplazos.items():
        if viejo.lower() in p.text.lower():
            for run in p.runs:
                if viejo.lower() in run.text.lower():
                    run.text = run.text.replace(viejo, nuevo)
                    cambiados += 1
            # Si no estaba en runs individuales, reemplazamos el párrafo entero
            if viejo.lower() in p.text.lower():
                p.text = p.text.replace(viejo, nuevo)
                cambiados += 1

# También en tablas (índice de figuras suele ser tabla)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for viejo, nuevo in reemplazos.items():
                    if viejo.lower() in p.text.lower():
                        p.text = p.text.replace(viejo, nuevo)
                        cambiados += 1

doc.save("INFORME_FINAL_SISTEMA.docx")
print(f"Renombrados {cambiados} títulos de diagramas.")

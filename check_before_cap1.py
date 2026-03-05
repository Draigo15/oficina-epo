# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA.docx")

for i, p in enumerate(doc.paragraphs):
    if "CAPÍTULO I" in p.text or "CAPITULO I" in p.text.upper():
        print(f"Cap I encontrado en par {i}")
        # Mostrar los 10 párrafos anteriores
        for j in range(max(0, i-10), i+1):
            t = doc.paragraphs[j].text
            style = doc.paragraphs[j].style.name
            print(f"  [{j}] [{style}] repr: {repr(t[:60])}")
        break

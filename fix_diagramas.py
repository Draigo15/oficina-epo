# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

def insert_placeholder(anchor_p, texto):
    """Inserta un parrafo marcador justo despues del ancla."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "BodyText")
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    i_elem = OxmlElement("w:i")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7F7F7F")
    rPr.append(i_elem)
    rPr.append(color)
    r.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.text = texto
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t_elem)
    p.append(r)
    anchor_p.addnext(p)
    return p

# -------------------------------------------------------
# 1. DIAGRAMA DE COMPONENTES [1091]: eliminar [1092..1859]
# -------------------------------------------------------
# Encontrar la posicion actual del titulo por texto
def find_by_text(keyword, style_hint=None):
    for i, p in enumerate(doc.paragraphs):
        if keyword.lower() in p.text.strip().lower():
            if style_hint is None or style_hint in p.style.name:
                return i, p
    return None, None

# Primero identificar todos los anclas ANTES de eliminar (los indices cambian)
# Usamos referencias a _p directamente

# Localizar por texto unico
anchors = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "Diagrama de componentes del sistema":
        anchors["componentes"] = (i, p._p)
    if t == "Diagrama de Despliegue" and i > 1000:
        anchors["despliegue"] = (i, p._p)

print("Anclas encontradas:")
for k, (i, _) in anchors.items():
    print(f"  {k}: [{i}]")

# --- Bloque componentes: encontrar rango a eliminar ---
comp_idx = anchors["componentes"][0]
comp_p = anchors["componentes"][1]

# Los parrafos a eliminar son desde comp_idx+1 hasta el siguiente "Nota. Elaboracion"
# sabemos que es hasta ~[1859] (antes de Despliegue)
# Buscamos el siguiente heading o "Diagrama de Despliegue"
end_comp = None
for i in range(comp_idx+1, len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    s = doc.paragraphs[i].style.name
    if t == "Diagrama de Despliegue" and i > comp_idx + 100:
        end_comp = i
        break

print(f"Rango componentes a eliminar: [{comp_idx+1}..{end_comp-1}] = {end_comp-comp_idx-1} parrafos")

# Eliminar de atras hacia adelante (usar lista de _p)
to_del = [doc.paragraphs[i]._p for i in range(comp_idx+1, end_comp)]
for xp in reversed(to_del):
    xp.getparent().remove(xp)

print(f"Componentes: eliminados {len(to_del)} parrafos")

# Insertar placeholder despues del titulo de componentes
insert_placeholder(comp_p, "[Insertar imagen: Diagrama de componentes del sistema]")

# --- Bloque despliegue: limpiar cuerpo vacio ---
# Buscar de nuevo por texto (indices cambiaron)
desp_p = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Diagrama de Despliegue" and i > 1000:
        desp_p = p._p
        desp_idx = i
        break

if desp_p:
    # El despliegue solo tenia parrafos vacios + "Nota. Elaboracion"
    # Insertar placeholder
    insert_placeholder(desp_p, "[Insertar imagen: Diagrama de Despliegue]")
    print(f"Despliegue: placeholder insertado en [{desp_idx}]")

# --- Diagramas de Secuencia: insertar placeholders donde esten vacios ---
seq_diags = [
    ("Diagrama de Secuencia CS-01", "[Insertar imagen: Diagrama de Secuencia CS-01 - Visualizar Dashboard]"),
    ("Diagrama de Secuencia CS-06", "[Insertar imagen: Diagrama de Secuencia CS-06 - Iniciar Sesion]"),
    ("Diagrama de Secuencia CS-07", "[Insertar imagen: Diagrama de Secuencia CS-07 - Administrar Dashboard Admin]"),
    ("Diagrama de Secuencia CS-02", "[Insertar imagen: Diagrama de Secuencia CS-02 - Gestionar Tareas]"),
    ("Diagrama de Secuencia CS-03", "[Insertar imagen: Diagrama de Secuencia CS-03 - Generar Reporte PDF]"),
]

added_seq = set()
for keyword, placeholder in seq_diags:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if keyword.lower() in t.lower() and i > 980 and keyword not in added_seq:
            # Verificar si el siguiente parrafo con texto ya es una Nota o Figura
            next_content = ""
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                nxt = doc.paragraphs[j].text.strip()
                if nxt:
                    next_content = nxt
                    break
            # Solo insertar si no hay contenido real (solo Nota o Figura)
            if not next_content or next_content.startswith("Nota.") or next_content.startswith("Figura"):
                insert_placeholder(p._p, placeholder)
                print(f"Secuencia: placeholder en [{i}] '{keyword[:40]}'")
                added_seq.add(keyword)
                break

# --- Diagrama de base de datos: ya esta correcto, no tocar ---

print(f"\nTotal parrafos final: {len(doc.paragraphs)}")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado -> guardado en {out}")

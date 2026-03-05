# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_Limpio.docx")

# Lista de palabras clave que sabemos que NO están en tu sistema actual
keywords_to_remove = [
    "Noticias", 
    "Programas Académicos", 
    "Información Institucional",
    "Mensaje de Contacto", 
    "Registrar estudiante",
    "Registrarse como Estudiante",
    "Prospectos", 
    "Leads",
    "Actualizar Contenido", 
    "Gestionar Contenido",
    "Mnesaje de Contacto",  # typo en el doc
    "Auntenticar"
]

# Casos de uso correctos:
# CS-01: Visualizar HOME/INICIO (o Dashboard)
# CS-02: Consultar Tareas Asignadas
# CS-06: Autenticar Usuario
# CS-07: Administrar Dashboard / Gestionar Notificaciones
# CS-08: Generar Reportes PDF

deleted_paras = 0

def has_invalid_keyword(text):
    text_lower = text.lower()
    for kw in keywords_to_remove:
        if kw.lower() in text_lower:
            return True
    return False

# Borramos los párrafos (índice, títulos, descripciones) que contengan los keywords inválidos
for p in doc.paragraphs:
    if has_invalid_keyword(p.text):
        p._element.getparent().remove(p._element)
        deleted_paras += 1

# Eliminar las tablas (Narrativas de Casos de Uso) que corresponden a los inválidos
deleted_tables = 0
for table in doc.tables:
    # Chequeamos la primera celda o texto de la tabla para ver sobre qué trata
    try:
        first_row_text = table.rows[0].cells[0].text + " " + table.rows[0].cells[1].text
        if has_invalid_keyword(first_row_text):
            table._element.getparent().remove(table._element)
            deleted_tables += 1
    except:
        pass

doc.save("INFORME_Actualizado_Sistema.docx")
print(f"Limpieza completada: {deleted_paras} párrafos y {deleted_tables} tablas eliminadas.")

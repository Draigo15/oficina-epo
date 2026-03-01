"""
Script v2 - Cambios restantes del INFORME.docx
  · Fechas de prácticas
  · 4.2.2 Problemas específicos
  · 4.3 Justificación (técnica, económica, administrativa, académica)
  · 4.3.2 Importancia
  · 4.4.2 Objetivos específicos
  · 4.5 Alcances
  · 4.6 Solución al Problema (descripción + componentes)
"""

import os
from docx import Document

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)

# ─── Utilidad base ─────────────────────────────────────────────────────────────
def set_para(para, text):
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ''
    if not para.runs:
        para.add_run(text)

def find_and_replace(fragment, new_text, limit=20, min_idx=0):
    """Busca el párrafo cuyo texto contenga 'fragment' y lo reemplaza."""
    key = fragment.lower()[:55]
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if i < min_idx:
            continue
        if key in para.text.lower() and count < limit:
            set_para(para, new_text)
            count += 1
            if limit == 1:
                print(f"    ✓ '{fragment[:60]}'")
                return True
    if count:
        print(f"    ✓ '{fragment[:60]}' ({count}x)")
    return count > 0

# ══════════════════════════════════════════════════════════════════════════════
# 1. FECHAS DE PRÁCTICAS
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- FECHAS ---")
find_and_replace("21 de marzo de 2025",   "15 de setiembre de 2025", limit=5)
find_and_replace("04 de julio de 2025",   "14 de diciembre de 2025", limit=5)
find_and_replace("21 de Marzo de 2025",   "15 de setiembre de 2025", limit=5)
find_and_replace("04 de Julio de 2025",   "14 de diciembre de 2025", limit=5)
# Texto párrafo de fechas en Cap I
find_and_replace(
    "Las actividades laborales se desarrollan en el área del Comité de mejora Continua de la Escuela Prof",
    "Las actividades laborales se desarrollan en el área del Comité de Mejora Continua (CMC) de la Escuela Profesional de Odontología en la Facultad de Ciencias de la Salud, Universidad Privada de Tacna.",
    limit=2
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. CAPÍTULO IV — 4.2.2 Problemas específicos
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.2.2 PROBLEMAS ESPECÍFICOS ---")

pe1 = ("¿Cómo afecta la gestión descentralizada de prospectos acadé-\n"
       "micos mediante registros dispersos en hojas de cálculo indepen-\n"
       "dientesalacapacidaddeseguimientosistemáticodeestudiantes\n"
       "potenciales y a la conversión de interesados en matriculados?")
pe1_nuevo = (
    "¿Cómo afecta la gestión manual y descentralizada de tareas administrativas "
    "mediante registros en papel y hojas de cálculo independientes a la capacidad "
    "del personal de la Oficina EPO para hacer seguimiento sistemático de "
    "actividades, cumplir plazos y generar evidencia documental de las labores "
    "realizadas mensualmente?"
)

pe2 = ("¿De qué manera impacta la ausencia de una plataforma centrali-\n"
       "zadadecomunicacióninstitucionalyladependenciatécnicapara")
pe2_nuevo = (
    "¿De qué manera impacta la ausencia de un sistema automatizado de generación "
    "de reportes en la capacidad del área para elaborar informes mensuales de "
    "productividad, presentar evidencias de actividades ante organismos de "
    "acreditación y tomar decisiones basadas en datos cuantitativos objetivos?"
)

pe3 = ("¿Cómo incide el proceso manual de inscripción de estudiantes")
pe3_nuevo = (
    "¿Cómo incide la falta de diferenciación de roles y permisos de usuario "
    "en la seguridad de la información administrativa, en la implementación de "
    "flujos de trabajo con responsabilidades definidas y en la auditoría de "
    "acciones realizadas sobre las tareas del área?"
)

pe4 = ("¿Dequéformalaausenciadeunsistemaespecializadoparages-")
pe4_nuevo = (
    "¿De qué forma la ausencia de un historial digital centralizado de tareas "
    "completadas dificulta el análisis de tendencias de productividad, la "
    "identificación de cargas de trabajo por usuario y la generación de "
    "estadísticas comparativas por período para la mejora continua del área?"
)

pe5 = ("¿De qué manera la falta de diferenciación de roles y permisos")
pe5_nuevo = (
    "¿De qué manera la inexistencia de un sistema de notificaciones internas "
    "afecta la comunicación oportuna entre el personal de la oficina respecto "
    "al estado de las tareas asignadas, las fechas de vencimiento próximas y "
    "los cambios en la prioridad de las actividades?"
)

for old, new in [(pe1, pe1_nuevo),(pe2, pe2_nuevo),(pe3, pe3_nuevo),(pe4, pe4_nuevo),(pe5, pe5_nuevo)]:
    find_and_replace(old[:55], new, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 3. JUSTIFICACIÓN TÉCNICA (4.3.1)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.3.1 JUSTIFICACIÓN TÉCNICA ---")

jt_p1_old = "La implementación del Sistema Integral de Gestión Aca- démica(SIGAOdontología)enlaEscuelaProfesionaldeOdon-"
jt_p1_nuevo = (
    "La implementación del Sistema de Gestión de Tareas – Oficina EPO en el "
    "Comité de Mejora Continua de la Escuela Profesional de Odontología de la "
    "Universidad Privada de Tacna representa una necesidad crítica para modernizar "
    "y digitalizar los procesos administrativos del área, generando beneficios "
    "tangibles que impactan positivamente en la eficiencia operativa del personal "
    "y en la calidad de la gestión institucional. La gestión manual de tareas "
    "mediante papeles, correos electrónicos dispersos y hojas de cálculo sin "
    "integración genera ineficiencias significativas que afectan la productividad "
    "del personal administrativo y la trazabilidad de actividades."
)

jt_p2_old = "nera ineficiencias significativas que afectan la productividad del personal administrativo y la expe"
jt_p2_nuevo = (
    "Un sistema web centralizado basado en arquitectura cliente-servidor con "
    "stack tecnológico MERN permitirá consolidar toda la información de tareas "
    "en una plataforma unificada accesible desde cualquier dispositivo con "
    "conexión a internet, reduciendo el tiempo de consulta del estado de "
    "actividades de varios minutos a consultas instantáneas. La arquitectura "
    "propuesta con API RESTful, autenticación JWT y roles diferenciados (Jefa "
    "y Asistente) garantiza seguridad en el manejo de información y proporciona "
    "escalabilidad para incorporar futuras funcionalidades sin rediseños costosos."
)

for old, new in [(jt_p1_old, jt_p1_nuevo),(jt_p2_old, jt_p2_nuevo)]:
    find_and_replace(old[:55], new, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 4. JUSTIFICACIÓN ECONÓMICA
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- JUSTIFICACIÓN ECONÓMICA ---")

je_old = "El sistema generará ahorros sustanciales al optimizar el uso del tiempo del personal administrativo"
je_nuevo = (
    "El sistema generará ahorros sustanciales al optimizar el uso del tiempo "
    "del personal administrativo, eliminando procesos manuales repetitivos que "
    "actualmente consumen horas laborales en tareas de bajo valor agregado. La "
    "automatización del registro de tareas, seguimiento del estado de actividades "
    "y generación de reportes mensuales en PDF liberará aproximadamente 3-5 horas "
    "semanales del personal, permitiendo reasignación de esfuerzos hacia "
    "actividades estratégicas de mayor impacto institucional. La reducción de "
    "errores mediante validación automática de datos disminuirá reprocesos "
    "administrativos que actualmente generan demoras. La generación automática "
    "de reportes PDF mensuales eliminará el tiempo invertido en consolidación "
    "manual de información dispersa en hojas de cálculo independientes, "
    "reduciendo el tiempo de elaboración de informes de varias horas a segundos."
)
find_and_replace(je_old[:55], je_nuevo, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 5. JUSTIFICACIÓN ADMINISTRATIVA
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- JUSTIFICACIÓN ADMINISTRATIVA ---")

ja_old = "La plataforma proporcionará control administrativo robus- to mediante centralización de información"
ja_nuevo = (
    "La plataforma proporcionará control administrativo robusto mediante "
    "centralización de información de tareas en base de datos unificada con "
    "trazabilidad completa de todas las operaciones realizadas por usuarios del "
    "sistema. El sistema de roles diferenciados (Jefa y Asistente) garantizará "
    "que cada usuario acceda únicamente a las funcionalidades apropiadas según "
    "su perfil: la Jefa podrá crear, editar, eliminar y reasignar tareas, mientras "
    "que el Asistente podrá registrar y completar las tareas asignadas. La "
    "digitalización completa del flujo de trabajo eliminará registros dispersos "
    "y documentos físicos, centralizando la información en un repositorio digital "
    "con historial completo. El dashboard ejecutivo con indicadores clave de "
    "desempeño proporcionará al área información actualizada en tiempo real "
    "sobre la carga de trabajo, tareas completadas y pendientes, facilitando la "
    "toma de decisiones basada en datos cuantitativos y el cumplimiento de "
    "requisitos de acreditación institucional mediante evidencia documentada."
)
find_and_replace(ja_old[:55], ja_nuevo, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 6. JUSTIFICACIÓN ACADÉMICA
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- JUSTIFICACIÓN ACADÉMICA ---")

jac_old = "La implementación del sistema contribuirá directamen- te a mejorar la calidad de servicios educativo"
jac_nuevo = (
    "La implementación del sistema contribuirá directamente a mejorar la calidad "
    "de la gestión administrativa del Comité de Mejora Continua mediante la "
    "digitalización y trazabilidad de sus procesos operativos. El personal "
    "administrativo dispondrá de una plataforma centralizada para registrar y "
    "consultar el estado de todas las tareas institucionales sin intermediarios "
    "ni búsquedas en archivos físicos, optimizando su tiempo y reduciendo la "
    "carga cognitiva asociada al seguimiento manual. El procesamiento automático "
    "de estadísticas de productividad proporcionará retroalimentación objetiva "
    "basada en datos verificables, facilitando la implementación de planes de "
    "mejora del área y fortaleciendo la cultura de mejora continua institucional. "
    "La generación sistemática de reportes mensuales permitirá identificar "
    "tendencias de carga de trabajo, evaluar la distribución de actividades "
    "entre el personal y diseñar estrategias de organización alineadas con los "
    "objetivos del CMC y los requerimientos de acreditación académica."
)
find_and_replace(jac_old[:55], jac_nuevo, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 7. IMPORTANCIA (4.3.2)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.3.2 IMPORTANCIA ---")

imp_old = "La modernización integral de la gestión académica y administra- tiva mediante tecnologías web modern"
imp_nuevo = (
    "La modernización integral de la gestión administrativa mediante tecnologías "
    "web modernas es estratégica para la transformación digital del Comité de "
    "Mejora Continua de la EPO. El sistema establecerá bases sólidas para futuras "
    "iniciativas de digitalización institucional, mejorará la imagen de eficiencia "
    "y transparencia del área, y posicionará al CMC como referente en adopción "
    "de tecnologías digitales aplicadas a la gestión administrativa dentro del "
    "ecosistema universitario. El uso de stack MERN (MongoDB, Express.js, React, "
    "Node.js), API RESTful, autenticación JWT, roles diferenciados y generación "
    "de reportes PDF con jsPDF fortalecerá el perfil técnico del área."
)
find_and_replace(imp_old[:55], imp_nuevo, limit=1)

imp2_old = "lidades adicionales (sistema de biblioteca virtual, plataforma de telemedi"
imp2_nuevo = (
    "El sistema permitirá incorporar futuras funcionalidades como integración "
    "con calendarios institucionales, módulo de asignación masiva de tareas, "
    "reportes por usuario y período, y sistema de notificaciones en tiempo real "
    "mediante extensión modular sin requerir rediseños arquitectónicos costosos. "
    "La separación entre frontend React y backend Node.js/Express permite "
    "incluso desarrollar aplicaciones móviles nativas que consuman la misma API, "
    "maximizando el retorno de inversión inicial en el desarrollo del sistema."
)
find_and_replace(imp2_old[:55], imp2_nuevo, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 8. OBJETIVOS ESPECÍFICOS (4.4.2)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.4.2 OBJETIVOS ESPECÍFICOS ---")

oe_bullets = [
    # (fragmento original, texto nuevo)
    (
        "Implementar un módulo centralizado de gestión de prospec-",
        "Implementar un módulo centralizado de gestión de tareas que permita "
        "crear, editar, eliminar y hacer seguimiento de actividades con título, "
        "descripción, prioridad diferenciada (normal/alta) y fecha límite, "
        "registrando automáticamente la fecha y hora de completado para garantizar "
        "trazabilidad completa del trabajo diario del área."
    ),
    (
        "Desarrollar una plataforma centralizada de comunicación ins-",
        "Desarrollar un módulo de generación automática de reportes mensuales "
        "en formato PDF con jsPDF y jsPDF-autotable que consolide las tareas "
        "completadas en un período seleccionado, eliminando la elaboración manual "
        "de informes y proporcionando evidencia documental estructurada de las "
        "actividades realizadas por el personal del CMC."
    ),
    (
        "Automatizar el proceso de inscripción de estudiantes median-",
        "Implementar un dashboard interactivo con estadísticas en tiempo real "
        "de productividad del área (tareas totales, completadas, pendientes, de "
        "alta prioridad) mediante gráficos visualizados con Recharts, facilitando "
        "el análisis rápido del estado del área y la toma de decisiones basada "
        "en indicadores cuantitativos objetivos."
    ),
    (
        "Implementar un sistema especializado de gestión de eventos",
        "Establecer un sistema de notificaciones internas que mantenga al "
        "personal informado sobre actualizaciones relevantes del estado de tareas, "
        "vencimientos próximos y cambios prioritarios, mejorando la comunicación "
        "interna del área sin depender de canales externos como correo electrónico "
        "o grupos de mensajería."
    ),
    (
        "Establecer un sistema de autenticación y autorización basa-",
        "Establecer un sistema de autenticación y autorización basado en roles "
        "diferenciados (Jefa y Asistente) mediante JSON Web Tokens (JWT) que "
        "garantice control de acceso granular a funcionalidades según perfil de "
        "usuario, asegure la información administrativa y proporcione trazabilidad "
        "de las acciones realizadas sobre las tareas del sistema."
    ),
]

for old_frag, new_text in oe_bullets:
    find_and_replace(old_frag[:55], new_text, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 9. ALCANCES (4.5)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.5 ALCANCES ---")

alc_p1_old = "El alcance del Sistema Integral de Gestión Académica (SIGA Odonto- logía) para la Escuela Profesiona"
alc_p1_nuevo = (
    "El alcance del Sistema de Gestión de Tareas – Oficina EPO para el Comité "
    "de Mejora Continua de la Escuela Profesional de Odontología de la Universidad "
    "Privada de Tacna comprende el análisis, diseño, desarrollo, pruebas e "
    "implementación de una aplicación web full-stack orientada a centralizar y "
    "automatizar la gestión de tareas administrativas del área. El sistema cubre, "
    "de forma integral, la creación y gestión de tareas con título, descripción, "
    "prioridad (normal/alta) y fecha límite; el seguimiento del estado de tareas "
    "(pendiente/completada) con registro automático de fecha y hora de completado; "
    "la generación automática de reportes mensuales en PDF; un dashboard con "
    "estadísticas de productividad en tiempo real; un sistema de notificaciones "
    "internas; y la administración de usuarios con autenticación JWT y "
    "autorización diferenciada por roles (Jefa y Asistente)."
)

alc_p2_old = "El sistema incluye autenticación segura mediante encriptación de con- traseñas con bcrypt"
alc_p2_nuevo = (
    "El sistema incluye autenticación segura mediante encriptación de contraseñas "
    "con bcrypt y generación de tokens JWT con expiración configurable, control "
    "de acceso por roles mediante middleware de verificación que restringe "
    "funcionalidades según perfil de usuario, interfaces de usuario responsive "
    "desarrolladas con React 18 y Tailwind CSS compatibles con dispositivos "
    "móviles y tablets, API RESTful con endpoints organizados por recursos "
    "(autenticación, tareas, reportes, notificaciones), generación de reportes "
    "PDF con jsPDF y jsPDF-autotable, visualización de estadísticas con Recharts, "
    "organización visual de tareas mediante arrastrar y soltar con React Beautiful "
    "DnD, manejo de fechas con date-fns y Moment.js, y manejo centralizado de "
    "errores con mensajes descriptivos."
)

alc_p3_old = "En su fase de desarrollo, el proyecto abarca la normalización y estruc- turacióndelmodelodedatosenMongoDBcon"
alc_p3_nuevo = (
    "En su fase de desarrollo, el proyecto abarca el modelado de datos en MongoDB "
    "con tres colecciones especializadas (usuarios, tareas, notificaciones), la "
    "implementación de interfaces de usuario diferenciadas para el rol Jefa "
    "(dashboard completo, CRUD de tareas, reportes, estadísticas, notificaciones) "
    "y el rol Asistente (vista de tareas asignadas, completado de actividades, "
    "notificaciones), la configuración del entorno de desarrollo con Vite para "
    "frontend y Node.js/Nodemon para backend, y la puesta en producción sobre "
    "infraestructura cloud mediante MongoDB Atlas para base de datos, Render "
    "para el backend y Vercel para el frontend."
)

alc_p4_old = "El alcance contempla también la definición de reglas de negocio para validación de datos académicos"
alc_p4_nuevo = (
    "El alcance contempla también la definición de reglas de negocio para "
    "validación de datos (campos requeridos, longitud máxima de textos, formato "
    "de fechas), la configuración de permisos diferenciados que determinan qué "
    "funcionalidades puede acceder cada rol, la generación de interfaces "
    "especializadas con filtros por estado de tarea (pendiente/completada), "
    "vistas de productividad mensual con selección de período, panel de "
    "notificaciones con marcado de leídas/no leídas, y perfil de usuario editable."
)

alc_p5_old = "Se incluye la capacitación funcional básica a usuarios clave (personal administrativoquegestionará"
alc_p5_nuevo = (
    "Se incluye la capacitación funcional básica a los usuarios del sistema "
    "(Jefa y Asistente del CMC), documentación técnica de arquitectura del "
    "sistema con diagramas de componentes y base de datos, manual de usuario "
    "con guías de uso por rol, lineamientos básicos de respaldo de base de datos "
    "MongoDB y el despliegue en ambiente de producción accesible desde navegadores "
    "web modernos sin requerir instalación de software adicional en los equipos."
)

alc_p6_old = "Quedan expresamente fuera del alcance la adquisición de servidores dedicad"
alc_p6_nuevo = (
    "Quedan expresamente fuera del alcance la adquisición de servidores dedicados "
    "on-premise (se utiliza infraestructura cloud existente), el desarrollo de "
    "aplicaciones móviles nativas para iOS o Android (el sistema web responsive "
    "cubre esta necesidad), la integración automática con sistemas legacy de la "
    "universidad central (SIGU, sistema de biblioteca), el desarrollo de módulos "
    "de gestión de notas o calificaciones académicas, la implementación de "
    "notificaciones push en tiempo real mediante WebSockets, y cambios en procesos "
    "organizacionales o normativas institucionales ajenas a los módulos funcionales "
    "definidos. Dichos elementos podrán evaluarse como ampliaciones futuras una "
    "vez que el sistema alcance estabilización operativa y se cuente con "
    "retroalimentación de los usuarios reales del área."
)

for old, new in [
    (alc_p1_old, alc_p1_nuevo),
    (alc_p2_old, alc_p2_nuevo),
    (alc_p3_old, alc_p3_nuevo),
    (alc_p4_old, alc_p4_nuevo),
    (alc_p5_old, alc_p5_nuevo),
    (alc_p6_old, alc_p6_nuevo),
]:
    find_and_replace(old[:55], new, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 10. SOLUCIÓN AL PROBLEMA (4.6)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.6 SOLUCIÓN AL PROBLEMA ---")

sol_intro_old = "Para abordar integralmente las deficiencias críticas identificadas en la gestión académica, administ"
sol_intro_nuevo = (
    "Para abordar integralmente las deficiencias identificadas en la gestión "
    "administrativa del Comité de Mejora Continua de la Escuela Profesional de "
    "Odontología, se propuso el desarrollo e implementación del Sistema de Gestión "
    "de Tareas – Oficina EPO, una plataforma web full-stack moderna basada en "
    "stack tecnológico MERN (MongoDB, Express.js, React y Node.js) con "
    "arquitectura cliente-servidor que centraliza y automatiza completamente el "
    "registro, seguimiento y reporte de tareas administrativas del área."
)

sol_desc_old = "El SIGA Odontología constituye una plataforma web completa y escalable para gestionar eficientemente"
sol_desc_nuevo = (
    "El Sistema de Gestión de Tareas – Oficina EPO constituye una plataforma "
    "web completa y escalable para gestionar eficientemente las tareas "
    "administrativas del CMC. Implementa una arquitectura de cliente-servidor "
    "con API RESTful que garantiza separación de responsabilidades, escalabilidad "
    "horizontal, mantenibilidad del código y flexibilidad tecnológica para "
    "adaptarse a futuras necesidades del área mediante extensión modular sin "
    "requerir rediseños arquitectónicos costosos."
)

sol_aborda_old = "La solución aborda los aspectos críticos de la problemática iden- tificada: centralización de gestió"
sol_aborda_nuevo = (
    "La solución aborda los aspectos críticos de la problemática identificada: "
    "centralización de tareas administrativas mediante módulo especializado con "
    "filtros por estado y prioridad; generación automática de reportes mensuales "
    "en PDF que eliminan la elaboración manual de informes; dashboard ejecutivo "
    "con estadísticas de productividad en tiempo real visualizadas con Recharts; "
    "sistema de notificaciones internas para comunicación oportuna entre el "
    "personal; y diferenciación de roles y permisos mediante autenticación JWT "
    "con control de acceso granular. La plataforma es accesible desde navegadores "
    "web modernos en computadoras de escritorio, laptops, tablets y dispositivos "
    "móviles mediante diseño responsive desarrollado con Tailwind CSS."
)

for old, new in [
    (sol_intro_old, sol_intro_nuevo),
    (sol_desc_old,  sol_desc_nuevo),
    (sol_aborda_old, sol_aborda_nuevo),
]:
    find_and_replace(old[:55], new, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 11. COMPONENTES PRINCIPALES (4.6.2)
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- 4.6.2 COMPONENTES PRINCIPALES ---")

comp_reemplazos = [
    (
        "Módulo de Gestión de Prospectos Académicos",
        "Módulo de Gestión de Tareas"
    ),
    (
        "Sistema especializado para registro centralizado de es- tudiantespotencialesconinformacióndecontactocompleta",
        "Sistema centralizado para creación y seguimiento de tareas con información completa "
        "(título, descripción, prioridad normal/alta, fecha límite, usuario asignado), "
        "registro automático de fecha y hora de completado con trazabilidad completa, "
        "filtros por estado (pendiente/completada) y prioridad, organización visual "
        "mediante arrastrar y soltar con React Beautiful DnD, y panel de gestión "
        "con acciones diferenciadas según rol (Jefa puede editar/eliminar, Asistente "
        "puede completar)."
    ),
    (
        "Módulo de Publicación de Noticias Institucionales",
        "Módulo de Reportes Mensuales en PDF"
    ),
    (
        "Plataforma de gestión de contenido editorial con editor de textoenriquecidopararedaccióndenoticias",
        "Sistema de generación automática de reportes mensuales en formato PDF "
        "mediante jsPDF y jsPDF-autotable, con selección de mes y año, listado "
        "tabular de tareas completadas en el período (título, descripción, prioridad, "
        "completado por, fecha de completado), resumen estadístico del período y "
        "exportación inmediata del documento sin procesamiento manual."
    ),
    (
        "Módulo de Gestión de Eventos Académicos",
        "Dashboard de Estadísticas en Tiempo Real"
    ),
    (
        "Sistema integral para programación de actividades for- mativas con registro de información completa",
        "Panel de control con indicadores clave de productividad del área: total de "
        "tareas registradas, tareas completadas, tareas pendientes y tareas de alta "
        "prioridad pendientes, visualizados mediante tarjetas de resumen y gráficos "
        "interactivos de barras y líneas con Recharts que muestran la evolución de "
        "productividad mensual de los últimos 6 meses."
    ),
    (
        "Módulo de Inscripción Digital de Estudiantes",
        "Sistema de Notificaciones Internas"
    ),
    (
        "Formulariowebcompletopararegistrodeaspirantescon validación automática de datos",
        "Sistema de notificaciones internas que informa al personal sobre "
        "actualizaciones relevantes del área: nuevas tareas asignadas, cambios "
        "de estado, tareas próximas a vencer y comunicados internos. Incluye "
        "panel de notificaciones con marcado de leídas/no leídas, contador de "
        "notificaciones no leídas en la barra de navegación y eliminación de "
        "notificaciones procesadas."
    ),
    (
        "Sistema de Gestión de Contenidos",
        "Sistema de Autenticación y Autorización por Roles"
    ),
    (
        "Interfaz administrativa intuitiva para actualización autó- noma de contenido del sitio web institucional",
        "Implementación de autenticación segura mediante JSON Web Tokens (JWT) "
        "con expiración configurable, encriptación de contraseñas con bcrypt "
        "(cost factor 10), diferenciación de roles con permisos granulares: "
        "Jefa con acceso completo (CRUD de tareas, reportes, estadísticas, "
        "notificaciones, gestión de usuarios) y Asistente con acceso restringido "
        "(vista de tareas, completado de actividades, notificaciones propias), "
        "middleware de verificación de roles que restringe el acceso a endpoints "
        "según perfil de usuario, y persistencia de sesión mediante token JWT "
        "almacenado en el cliente."
    ),
    (
        "Módulo de Mensajes de Contacto",
        "Módulo de Perfil de Usuario"
    ),
    (
        "Formulario público de contacto accesible desde el sitio web con validación de campos obligatorios",
        "Sección de perfil personal que permite a cada usuario visualizar y "
        "actualizar su información básica (nombre completo, nombre de usuario), "
        "proporciona resumen de actividad (tareas creadas, completadas) y "
        "permite cambio de contraseña con validación de seguridad, garantizando "
        "que cada usuario mantenga control sobre su cuenta de acceso al sistema."
    ),
]

for old_frag, new_text in comp_reemplazos:
    find_and_replace(old_frag[:55], new_text, limit=1)

# ══════════════════════════════════════════════════════════════════════════════
# 12. ELIMINAR referencias residuales al sistema de Odontología
# ══════════════════════════════════════════════════════════════════════════════
print("\n--- LIMPIEZA RESIDUAL ---")

residuales = [
    ("odontología", "la Oficina EPO"),
    ("Odontología", "la Oficina EPO"),
    ("SIGA Odontología", "Sistema de Gestión de Tareas – Oficina EPO"),
    ("Sistema Integral de Gestión Académica", "Sistema de Gestión de Tareas – Oficina EPO"),
    ("gestión académica", "gestión de tareas"),
    ("prospectos académicos", "tareas administrativas"),
    ("leads", "tareas"),
    ("noticias institucionales", "reportes mensuales"),
    ("eventos académicos", "actividades del área"),
    ("inscripciones estudiantiles", "registro de tareas"),
    ("estudiantes potenciales", "personal del área"),
    ("administrador, docente, estudiante", "Jefa y Asistente"),
    ("administrador, docente, estudiante", "Jefa y Asistente"),
    ("más de 40 endpoints", "endpoints organizados por recursos"),
]

# Solo aplicar en párrafos que TODAVÍA contengan estas palabras
# (evitando sobreescribir párrafos ya corregidos)
for old_frag, new_text in residuales:
    key = old_frag.lower()
    for para in doc.paragraphs:
        if key in para.text.lower():
            # Solo si el párrafo aún tiene terminología de Odontología no corregida
            full = para.text
            new_full = full.replace(old_frag, new_text)
            if new_full != full:
                for i, run in enumerate(para.runs):
                    if old_frag in run.text:
                        run.text = run.text.replace(old_frag, new_text)

print("    ✓ Limpieza de residuos completada")

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
doc.save(DEST)
print("\n✅ INFORME.docx v2 guardado con todos los cambios.")

# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

print(f"Total parrafos: {len(paras)}")

# 1. Encontrar el bloque mal insertado:
# Empieza en [186] con "Narrativas de Casos de Uso" hasta antes del RF de Cap II [438]
# Verificar los limites
inicio_mal = None
fin_mal = None

for i, p in enumerate(paras):
    t = p.text.strip()
    if t == "Narrativas de Casos de Uso" and p.style.name == "Heading 3" and i < 300:
        inicio_mal = i
    if "Requerimientos Funcionales" in t and p.style.name == "Heading 3" and inicio_mal and i > inicio_mal and fin_mal is None:
        fin_mal = i - 1  # el parrafo justo antes del RF de Cap II

print(f"Bloque mal insertado: [{inicio_mal}] - [{fin_mal}]")
print(f"  [{inicio_mal}]: '{paras[inicio_mal].text[:60]}'")
print(f"  [{fin_mal}]:   '{paras[fin_mal].text[:60]}'")
print(f"  [{fin_mal+1}]: '{paras[fin_mal+1].text[:60]}'")

# 2. Eliminar [inicio_mal..fin_mal]
count = 0
for idx in range(fin_mal, inicio_mal - 1, -1):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)
    count += 1
print(f"\nEliminados: {count} parrafos del bloque mal insertado.")

# 3. Verificar la nueva posicion del RF de Cap IV
print(f"\nTotal tras limpieza: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "Requerimientos Funcionales" in t and p.style.name == "Heading 3":
        print(f"[{i}] [{p.style.name}] '{t[:70]}'")
        print(f"  Previo: [{i-2}] [{doc.paragraphs[i-2].style.name}] '{doc.paragraphs[i-2].text[:60]}'")
        print(f"  Sig: [{i+1}] [{doc.paragraphs[i+1].style.name}] '{doc.paragraphs[i+1].text[:60]}'")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"\nGuardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

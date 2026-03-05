# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document("INFORME_FINAL_SISTEMA.docx")

def insert_paragraph_after(ref_para, text, style="Normal"):
    """Inserta un párrafo nuevo después del párrafo de referencia."""
    new_para = OxmlElement("w:p")
    ref_para._p.addnext(new_para)
    new_p_obj = doc.paragraphs[doc.paragraphs.index(ref_para) + 1]
    new_p_obj.style = doc.styles[style]
    new_p_obj.text = text
    return new_p_obj

# Buscar los párrafos vacíos que necesitan contenido
paragraphs = doc.paragraphs

for i, p in enumerate(paragraphs):
    text = p.text.strip()

    # === 2.6 UML: Diagrama de Componentes (ya tiene texto, verificar) ===
    if text == "UML: Diagrama de Componentes":
        # Revisar si el siguiente párrafo ya tiene contenido
        next_text = paragraphs[i+1].text.strip() if i+1 < len(paragraphs) else ""
        if next_text.startswith("UML:") or not next_text or len(next_text) < 20:
            insert_paragraph_after(p,
                "El Diagrama de Componentes UML representa las unidades de software del Sistema de Gestión de Tareas  Oficina EPO y sus dependencias. Incluye el frontend React (componentes, páginas, contextos y utilidades), el backend Node.js/Express (rutas, middleware, modelos Mongoose) y la base de datos MongoDB Atlas. Las interfaces provistas y requeridas entre componentes están definidas mediante la API RESTful con endpoints organizados por recursos (auth, tasks, reports, notifications).",
                "Normal")
        break

for i, p in enumerate(paragraphs):
    text = p.text.strip()

    # === 2.7 UML: Diagrama de Casos de Uso ===
    if text == "UML: Diagrama de Casos de Uso":
        next_text = paragraphs[i+1].text.strip() if i+1 < len(paragraphs) else ""
        if "Diagrama de Procesos" in next_text or len(next_text) < 20:
            insert_paragraph_after(p,
                "El Diagrama de Casos de Uso describe las interacciones entre los actores del sistema y sus funcionalidades. Los actores son: Jefa del CMC (rol administrador) y Asistente del CMC (rol usuario). Los casos de uso principales son: CS-01 Visualizar Dashboard, CS-02 Gestionar Tareas, CS-03 Generar Reporte PDF, CS-04 Gestionar Notificaciones, CS-05 Ver Estadísticas, CS-06 Autenticar Usuario y CS-07 Gestionar Perfil de Usuario. Cada caso de uso está respaldado por un endpoint de la API RESTful del backend.",
                "Normal")
        break

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()

    # === 2.8 Diagramas de Procesos ===
    if text == "Diagramas de Procesos":
        next_text = doc.paragraphs[i+1].text.strip() if i+1 < len(doc.paragraphs) else ""
        if "Tecnologías" in next_text or len(next_text) < 20:
            insert_paragraph_after(p,
                "Los Diagramas de Secuencia UML documentan el flujo temporal de mensajes entre los actores y el sistema. Se modelan los procesos principales: inicio de sesión con autenticación JWT, creación y actualización de tareas con cambio de estado Kanban, generación automática de reportes PDF con jsPDF, y el flujo de notificaciones automáticas para tareas vencidas. Estos diagramas evidencian la separación de responsabilidades entre el cliente React y el servidor Node.js/Express conectado a MongoDB Atlas.",
                "Normal")
        break

doc.save("INFORME_FINAL_SISTEMA.docx")
print("Secciones restauradas correctamente.")

# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA.docx")

# 1. Corregir estilo de los párrafos restaurados (Normal -> Body Text)
for p in doc.paragraphs:
    if ("Diagrama de Casos de Uso describe las interacciones" in p.text or
        "Diagramas de Secuencia UML documentan el flujo temporal" in p.text):
        p.style = doc.styles["Body Text"]
        print(f"  Estilo corregido a Body Text: {p.text[:60]}")

# 2. Eliminar la página en blanco antes de CAPÍTULO I
# Son los párrafos vacíos en indices 144 y 145 (Body Text vacios)
to_delete = []
for i, p in enumerate(doc.paragraphs):
    if "CAPÍTULO I" in p.text or "CAPITULO I" in p.text.upper():
        # Eliminar los párrafos vacíos justo antes
        j = i - 1
        while j >= 0 and doc.paragraphs[j].text.strip() == "":
            to_delete.append(doc.paragraphs[j])
            j -= 1
        break

for p in to_delete:
    p._element.getparent().remove(p._element)
    print(f"  Párrafo vacío eliminado antes de CAPÍTULO I")

doc.save("INFORME_FINAL_SISTEMA.docx")
print("Listo.")

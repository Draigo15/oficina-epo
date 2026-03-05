# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

fixes = 0
to_delete = []

for i, p in enumerate(paras):
    t = p.text.strip()
    if i > 138:
        break

    # --- CORRECCIONES DE TEXTO ---

    # CS-02 Consultar Tareas Asignadas -> Gestionar Tareas
    if "CS-02" in t and "Consultar Tareas Asignadas" in t:
        for run in p.runs:
            if "Consultar Tareas Asignadas" in run.text:
                run.text = run.text.replace("Consultar Tareas Asignadas", "Gestionar Tareas")
                fixes += 1
                print(f"[{i}] CS-02: Consultar -> Gestionar Tareas")

    # CS-06 Autenticar Usuario Administrativo -> Autenticar Usuario
    if "CS-06" in t and "Administrativo" in t:
        for run in p.runs:
            if "Administrativo" in run.text:
                run.text = run.text.replace("Autenticar Usuario Administrativo", "Autenticar Usuario")
                fixes += 1
                print(f"[{i}] CS-06: eliminado 'Administrativo'")

    # Diagrama Secuencia CS-01 Home/Inicio -> Dashboard
    if "CS-01" in t and ("Home/Inicio" in t or "Visualizar Home" in t):
        for run in p.runs:
            if "Home/Inicio" in run.text:
                run.text = run.text.replace("Visualizar Home/Inicio", "Visualizar Dashboard")
                run.text = run.text.replace("Home/Inicio", "Dashboard")
                fixes += 1
                print(f"[{i}] Secuencia CS-01: Home/Inicio -> Dashboard")

    # Typo con caracter extrano en "Nosotros"
    if "\u00e7on" in t or "Nosotros" in t and "institucional" in t:
        for run in p.runs:
            if "\u00e7on" in run.text:
                run.text = run.text.replace("\u00e7on", "con")
                fixes += 1
                print(f"[{i}] Typo corregido: con")

    # --- ENTRADAS DEL PORTAL A ELIMINAR (indice de figuras) ---
    portal_keywords = [
        "tipo de evento",          # filtros eventos
        "Nosotros",                # pagina Nosotros portal
        "Misi\u00f3n y Visi\u00f3n de la Escuela Profesional de Odontolog\u00eda",
        "informaci\u00f3n adicional sobre la instituci\u00f3n",
        "confirmaci\u00f3n de inscripci\u00f3n exitosa a evento",
        "lead manualmente",        # leads portal
        "lead con opciones",       # leads portal
        "reportes del sitio web",  # CMS portal
        "Editor de contenido",     # CMS portal
        "Opciones de configuraci\u00f3n de contenido",
        "Gestiones visuales de la p\u00e1gina",
        "elementos visuales de la p\u00e1gina",
        "configuraci\u00f3n avanzada de contenido",
        "im\u00e1genes y recursos multimedia",
        "inscripci\u00f3n exitosa a evento",
        "inicio de sesi\u00f3n como docente",
        "Dashboard del docente",
        "evento gestionado por el docente",
        "asistencia con registro de hora",
        "creaci\u00f3n de nuevo evento",
        "eventos creados por el docente",
        "Perfil del docente",
        "mas Acad\u00e9micos",  # fragmento roto
    ]
    for kw in portal_keywords:
        if kw.lower() in t.lower() and p.style.name == "List Paragraph":
            to_delete.append(i)
            print(f"[{i}] ELIMINAR (portal): '{t[:80]}'")
            break

print()
print(f"Entradas a eliminar: {len(to_delete)}")
for idx in sorted(to_delete, reverse=True):
    doc.paragraphs[idx]._p.getparent().remove(doc.paragraphs[idx]._p)

print(f"Correcciones de texto: {fixes}")
print(f"Entradas eliminadas: {len(to_delete)}")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"\nGuardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

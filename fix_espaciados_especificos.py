# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

# Patron especifico de texto espaciado REAL:
# letra suelta (1-2 chars) + espacio + fragmento, al menos 4 veces en el texto
def es_espaciado(text):
    matches = re.findall(r'\b[A-Za-záéíóúñÁÉÍÓÚÑ]{1,2}\s(?=[a-záéíóúñ])', text)
    return len(matches) >= 4

to_delete = []

for i, p in enumerate(paras):
    t = p.text.strip()
    if not t:
        continue
    # Solo en rangos especificos del portal viejo + sus estilos normales
    in_zona_portal_flujo = 410 <= i <= 580  # flujo inscripcion portal
    in_zona_bd_portal = 2790 <= i <= 2900   # diagrama BD portal

    if (in_zona_portal_flujo or in_zona_bd_portal) and es_espaciado(t):
        to_delete.append((i, t[:70]))

print(f"Parrafos a eliminar: {len(to_delete)}")
for idx, text in to_delete:
    print(f"  [{idx}] {text}")

# Eliminar de atras hacia adelante
for idx, text in sorted(to_delete, key=lambda x: x[0], reverse=True):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)

print(f"\n{len(to_delete)} parrafos eliminados.")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

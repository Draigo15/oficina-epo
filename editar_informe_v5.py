"""
Script v5 — Corrección de los 20 residuales detectados en auditoría completa
"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)
p    = doc.paragraphs

def ft(para):
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))

def set_xml(para, new_text):
    all_t = para._p.findall(f'.//{{{NS}}}t')
    if all_t:
        all_t[0].text = new_text
        for t in all_t[1:]: t.text = ''
    else:
        para.add_run(new_text)

def rep(para, old, new):
    full = ft(para)
    if old in full:
        set_xml(para, full.replace(old, new, 1))
        return True
    return False

# ─── [215] Roles del diagrama de casos de uso ────────────────────────────────
rep(p[215],
    'administrador, docente, estu- dian- te, visitante',
    'Jefa, Asistente')
rep(p[215],
    'administrador, docente, estudiante, visitante',
    'Jefa, Asistente')
rep(p[215],
    'administrador, docente, estu-\ndian- te',
    'Jefa, Asistente')
# Reemplazar subtexto más amplio si no encontró
full = ft(p[215])
if 'administrador, docente' in full:
    import re
    new = re.sub(r'administrador,\s*docente,\s*estu[- \n]*dian[- \n]*te,?\s*(visitante)?', 'Jefa, Asistente', full, flags=re.IGNORECASE)
    set_xml(p[215], new)
print(f'[215] → {ft(p[215])[:90]}')

# ─── [228] express-validator en Marco Teórico ────────────────────────────────
rep(p[228],
    'express-validator',
    'validación con Mongoose')
print(f'[228] → {ft(p[228])[:90]}')

# ─── [344] "gestión de eventos" en intro/cap ─────────────────────────────────
rep(p[344], 'gestión de eventos', 'gestión de tareas')
print(f'[344] → {ft(p[344])[:90]}')

# ─── [429] "gestión de eventos" en descripción ───────────────────────────────
rep(p[429], 'gestión de eventos', 'gestión de tareas del área')
print(f'[429] → {ft(p[429])[:90]}')

# ─── [1194] express-validator en caso de uso (flujo alternativo) ─────────────
rep(p[1194],
    '25c. Express-validator detecta datos inválidos no capturados en frontend → 26c. El backend devuelve',
    '25c. Validación Mongoose detecta datos inválidos → 26c. El backend devuelve')
rep(p[1194],
    'Express-validator detecta',
    'Validación Mongoose detecta')
print(f'[1194] → {ft(p[1194])[:90]}')

# ─── [1213] "registrations" en descripción CU ────────────────────────────────
rep(p[1213],
    'formulario de ins- cripción digital integral',
    'formulario de registro de tarea')
rep(p[1213], 'registrations', 'tareas')
print(f'[1213] → {ft(p[1213])[:90]}')

# ─── [1216] precondición con "registrations" y "tareas" ──────────────────────
rep(p[1216],
    "Las colecciones 'registrations' y 'tareas' deben existir en MongoDB.",
    "La colección 'tareas' debe existir en MongoDB.")
rep(p[1216], 'registrations', 'tareas')
print(f'[1216] → {ft(p[1216])[:90]}')

# ─── [1337] "express-validator" en test ──────────────────────────────────────
rep(p[1337],
    'El backend valida datos con express-validator.',
    'El backend valida datos con Mongoose schema validators.')
print(f'[1337] → {ft(p[1337])[:90]}')

# ─── [1338] "registrations" en precondición test ─────────────────────────────
rep(p[1338],
    "El sistema verifica que DNI no exista en 'registrations'.",
    "El sistema verifica que la tarea no esté duplicada en 'tareas'.")
rep(p[1338], 'registrations', 'tareas')
print(f'[1338] → {ft(p[1338])[:90]}')

# ─── [1339] "registrations" en acción test ───────────────────────────────────
rep(p[1339],
    "El sistema crea documento en 'registrations' con: datos comple- tos, estado: 'pendiente', createdAt.",
    "El sistema crea documento en 'tareas' con: título, descripción, prioridad, estado 'pendiente', createdAt.")
rep(p[1339],
    "El sistema crea documento en 'registrations'",
    "El sistema crea documento en 'tareas'")
rep(p[1339], 'registrations', 'tareas')
print(f'[1339] → {ft(p[1339])[:90]}')

# ─── [1352] "registrations" en resultado test ────────────────────────────────
rep(p[1352],
    "Se crea documento en 'registrations' con estado 'pendien- te' y todos los datos.",
    "Se crea documento en 'tareas' con estado 'pendiente', prioridad y responsable asignados.")
rep(p[1352], 'registrations', 'tareas')
print(f'[1352] → {ft(p[1352])[:90]}')

# ─── [1396] express-validator en flujo alternativo ───────────────────────────
rep(p[1396],
    '47d. Express-validator detecta campo inválido → 48d. El bac- kend devuelve código HTTP 400 con error',
    '47d. Mongoose detecta campo inválido → 48d. El backend devuelve código HTTP 400 con error')
rep(p[1396],
    'Express-validator detecta',
    'Mongoose detecta')
print(f'[1396] → {ft(p[1396])[:90]}')

# ─── [1648] express-validator en test ────────────────────────────────────────
rep(p[1648],
    'El backend valida datos con express-validator.',
    'El backend valida datos con Mongoose schema validators.')
print(f'[1648] → {ft(p[1648])[:90]}')

# ─── [1979] express-validator en test ────────────────────────────────────────
rep(p[1979],
    'El backend valida datos con express-validator.',
    'El backend valida datos con Mongoose schema validators.')
print(f'[1979] → {ft(p[1979])[:90]}')

# ─── [3414] /api/registrations endpoint ──────────────────────────────────────
rep(p[3414], '/api/registrations', '/api/notificaciones')
print(f'[3414] → {ft(p[3414])[:90]}')

# ─── [3480] registrationRoutes.js en árbol de archivos ───────────────────────
rep(p[3480],
    'registrationRoutes.js  POST /api/registrations',
    'notificationsRoutes.js  GET/POST /api/notificaciones')
rep(p[3480], 'registrationRoutes.js', 'notificationsRoutes.js')
rep(p[3480], '/api/registrations', '/api/notificaciones')
print(f'[3480] → {ft(p[3480])[:90]}')

# ─── [3481] GET/PUT /api/registrations ───────────────────────────────────────
rep(p[3481],
    'GET /api/registrations PUT /api/registrations/:id',
    'PUT /api/notificaciones/:id/read  (marcar leída)')
rep(p[3481], '/api/registrations', '/api/notificaciones')
print(f'[3481] → {ft(p[3481])[:90]}')

# ─── [3566] getUserRegistrations() ───────────────────────────────────────────
rep(p[3566],
    'getUserRegistrations()',
    'getUserNotifications()')
print(f'[3566] → {ft(p[3566])[:90]}')

# ─── [3760] Diagrama de arquitectura con "registrations" ─────────────────────
rep(p[3760], 'registrations', 'notificaciones')
print(f'[3760] → {ft(p[3760])[:110]}')

# ─── [4211] Nota beneficios con "gestión de eventos" ─────────────────────────
rep(p[4211],
    'publicación de contenido, seguimiento de tareas',
    'gestión y seguimiento de tareas')
rep(p[4211], 'gestión de eventos', 'gestión de tareas del área')
print(f'[4211] → {ft(p[4211])[:110]}')

# ─── Búsqueda global para express-validator y registrations restantes ─────────
print('\n--- Búsqueda global residual ---')
for i, para in enumerate(p):
    full = ft(para).lower()
    changed = False
    for old, new in [
        ('express-validator', 'Mongoose validators'),
        ('express validator',  'Mongoose validators'),
    ]:
        if old in full:
            rep(para, old, new)
            rep(para, old.capitalize(), new)
            print(f'  [{i}] express-validator → Mongoose validators')
            changed = True
            break
    full = ft(para).lower()
    for old, new in [
        ("'registrations'", "'tareas'"),
        ('/api/registrations', '/api/notificaciones'),
        ('registrationRoutes', 'notificationsRoutes'),
        ('getUserRegistrations', 'getUserNotifications'),
        ('gestión de eventos', 'gestión de tareas'),
        ('Gestión de Eventos', 'Gestión de Tareas'),
        ('administrador, docente, estudiante', 'Jefa, Asistente'),
    ]:
        if old.lower() in full:
            rep(para, old, new)
            print(f'  [{i}] "{old[:40]}" → "{new[:35]}"')
            full = ft(para).lower()

doc.save(DEST)
print('\n✅ v5 guardado.')

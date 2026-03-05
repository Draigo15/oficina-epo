# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

fixes = 0

# 1. Renombrar heading [4360] "Gestion Multimedia" -> "Gestion de Usuarios"
p = paras[4360]
if "Gesti" in p.text and "Multimedia" in p.text:
    for run in p.runs:
        if "Multimedia" in run.text:
            run.text = run.text.replace("Gestión Multimedia", "Gestión de Usuarios")
            fixes += 1
    if fixes == 0:
        # Si el texto esta distribuido en runs, limpiar y reescribir
        for run in p.runs:
            run.text = ""
        p.runs[0].text = "Gestión de Usuarios" if p.runs else None
        if p.runs:
            fixes += 1
    print(f"[4360] Renombrado a 'Gestión de Usuarios'")

# 2. Eliminar headings vacios [4368, 4369, 4383, 4384] (indices ajustados por fix 1)
# Buscar headings Heading 3 completamente vacios en zona 4360+
to_delete_empty = []
for i in range(4360, min(4395, len(paras))):
    p = paras[i]
    t = p.text.strip()
    if p.style.name == "Heading 3" and not t:
        to_delete_empty.append(i)

print(f"Headings vacios a eliminar: {to_delete_empty}")
for idx in sorted(to_delete_empty, reverse=True):
    p = doc.paragraphs[idx]
    p._p.getparent().remove(p._p)
    fixes += 1
    print(f"  [{idx}] Heading 3 vacio eliminado")

print(f"\nTotal fixes: {fixes}")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

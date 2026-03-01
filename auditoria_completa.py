import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc  = Document(DEST)
p    = doc.paragraphs

def ft(para):
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))

# ─── Términos del sistema del amigo que NO deben aparecer ─────────────────────
bad = [
    # Tecnologías no usadas
    ('multer',           'tecnología no usada'),
    ('cloudinary',       'tecnología no usada'),
    ('nodemailer',       'tecnología no usada'),
    ('docker',           'tecnología no usada'),
    ('nginx',            'tecnología no usada'),
    ('openapi 3',        'tecnología no usada'),
    ('rhf + zod',        'tecnología no usada'),
    ('react hook form',  'tecnología no usada'),
    ('express-validator','tecnología no usada'),
    ('concurrently',     'tecnología no usada'),
    ('pm2',              'tecnología no usada (process manager VPS)'),
    ('ubuntu server',    'infraestructura no usada'),
    ('aws s3',           'almacenamiento no usado'),
    ('sendgrid',         'email no usado'),
    ('oauth 2',          'auth no usada'),
    ('firebase',         'no usada'),
    # Sistema del amigo
    ('siga odontolog',   'nombre incorrecto del sistema'),
    ('siga la oficina',  'nombre incorrecto'),
    ('del siga',         'nombre incorrecto'),
    ('el siga',          'nombre incorrecto'),
    ('sistema integral de gestión académica', 'nombre incorrecto'),
    # Módulos del amigo
    ('prospectos académicos',  'módulo incorrecto'),
    ('gestión de prospectos',  'módulo incorrecto'),
    ('captación de leads',     'módulo incorrecto'),
    ('gestión de leads',       'módulo incorrecto'),
    ('leads académicos',       'módulo incorrecto'),
    ('noticias institucionales','módulo incorrecto'),
    ('noticias académicas',    'módulo incorrecto'),
    ('portal web institucional','módulo incorrecto'),
    ('portal público',         'módulo incorrecto'),
    ('eventos académicos',     'módulo incorrecto'),
    ('inscripción de estudiantes','módulo incorrecto'),
    ('módulo del estudiante',  'módulo incorrecto'),
    ('módulo del docente',     'módulo incorrecto'),
    ('gestión de noticias',    'módulo incorrecto'),
    ('gestión de eventos',     'módulo incorrecto'),
    # Roles incorrectos
    ('rol de estudiante',      'rol incorrecto'),
    ('rol de docente',         'rol incorrecto'),
    ('administrador, docente', 'roles incorrectos'),
    # Costos/fechas viejos
    ('16 semanas',  'duración incorrecta'),
    ('320 horas',   'horas incorrectas'),
    ('2,720',       'costo incorrecto'),
    ('3,020',       'costo incorrecto'),
    # Infraestructura incorrecta
    ('servidor vps', 'infraestructura incorrecta'),
    ('vps básico',   'infraestructura incorrecta'),
    ('vps basico',   'infraestructura incorrecta'),
    # Colecciones no existentes
    ('registrations', 'colección no existente'),
    ('/server/uploads','directorio incorrecto (multer)'),
]

print('=== AUDITORÍA COMPLETA ===\n')
issues = []

for i, para in enumerate(p):
    txt = ft(para).lower()
    for term, reason in bad:
        if term in txt:
            issues.append((i, term, reason, ft(para)[:100]))
            break  # solo reportar el primer match por párrafo

for table in doc.tables:
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.lower()
            for term, reason in bad:
                if term in txt:
                    issues.append(('tabla', term, reason, cell.text[:80]))
                    break

if issues:
    print(f'⚠ {len(issues)} problema(s) encontrado(s):\n')
    for loc, term, reason, preview in issues:
        print(f'  [{loc}] ({reason})')
        print(f'         término: "{term}"')
        print(f'         texto:   {preview}')
        print()
else:
    print('✅ Sin problemas tecnológicos/terminológicos detectados')

# ─── Verificar secciones clave que deben existir ─────────────────────────────
print('\n=== VERIFICACIÓN DE SECCIONES CLAVE ===\n')
must_have = [
    'Sistema de Gestión de Tareas',
    'jsPDF',
    'Recharts',
    'React Beautiful DnD',
    'date-fns',
    'MongoDB Atlas',
    'Render',
    'Vercel',
    'JWT',
    'bcrypt',
    'Jefa',
    'Asistente',
    '13 semanas',
    '2,210',
    '260 horas',
    'TIR',
    'VAN',
]
full_doc = '\n'.join(ft(para) for para in p)
for term in must_have:
    count = full_doc.count(term)
    if count == 0:
        print(f'  ✗ AUSENTE: "{term}"')
    else:
        print(f'  ✓ "{term}" ({count}x)')

print(f'\nTotal párrafos: {len(p)} | Tablas: {len(doc.tables)}')

import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document(r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx')

def ft(para):
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))

bad = [
    'multer', 'cloudinary', 'nodemailer', 'docker', 'nginx', 'openapi 3',
    'rhf + zod', 'react hook form',
    'siga odontolog', 'siga la oficina',
    '16 semanas', '2,720', '3,020',
    'vps básico', 'vps basico', 'servidor vps',
    'noticias académicas', 'noticias institucionales',
    'inscripción de estudiantes', 'inscripcion de estudiantes',
    'prospectos académicos', 'prospectos academicos',
    'captación de leads', 'gestión de leads',
    'módulo del estudiante', 'modulo del estudiante',
    'módulo del docente', 'modulo del docente',
    'portal público', 'portal web institucional',
    'events académicos', 'eventos académicos',
    'estudiantes potenciales',
    'rol de estudiante', 'rol de docente',
    'gestión de noticias', 'módulo de noticias',
]

issues = []
for i, para in enumerate(doc.paragraphs):
    txt = ft(para).lower()
    for b in bad:
        if b in txt:
            issues.append(f'P[{i}]: "{b}" → {ft(para)[:90]}')
            break

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.lower()
            for b in bad:
                if b in txt:
                    issues.append(f'T{ti}R{ri}C{ci}: "{b}" → {cell.text[:70]}')
                    break

if issues:
    print(f'PROBLEMAS DETECTADOS ({len(issues)}):')
    for iss in issues:
        print(f'  {iss}')
else:
    print('✅ AUDITORÍA FINAL OK – Sin residuales detectados')

print(f'\nTotal párrafos: {len(doc.paragraphs)} | Tablas: {len(doc.tables)}')

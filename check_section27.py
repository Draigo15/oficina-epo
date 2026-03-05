# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA.docx")

sections_to_check = ["UML: Diagrama de Componentes", "UML: Diagrama de Casos de Uso", "Diagramas de Procesos"]

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for sec in sections_to_check:
        if sec in text:
            print(f"\n=== SECCIÓN: {text} (par {i}) ===")
            # Mostrar los 10 párrafos siguientes
            for j in range(1, 12):
                if i+j < len(doc.paragraphs):
                    next_p = doc.paragraphs[i+j].text.strip()
                    if next_p:
                        print(f"  [{i+j}]: {next_p[:100]}")
            break

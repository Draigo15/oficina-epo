"""
Script v3c — Corrección de residuales con texto fragmentado en múltiples runs
             Usa manipulación XML directa
"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from lxml import etree

DEST = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
doc = Document(DEST)
paras = doc.paragraphs
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def para_full_text(para):
    """Concatena todos los w:t del párrafo incluyendo dentro de hyperlinks."""
    return ''.join(t.text or '' for t in para._p.findall(f'.//{{{NS}}}t'))


def replace_in_para_xml(para, old, new):
    """
    Substituye `old` por `new` a nivel XML aunque el texto esté fragmentado.
    Pone el texto resultante en el primer w:t y vacía los demás.
    """
    all_t = para._p.findall(f'.//{{{NS}}}t')
    full = ''.join(t.text or '' for t in all_t)
    if old not in full:
        return False
    replaced = full.replace(old, new, 1)
    if all_t:
        all_t[0].text = replaced
        for t in all_t[1:]:
            t.text = ''
    return True


fixes = [
    # (índice_párrafo, texto_viejo, texto_nuevo)

    # ── ÍNDICE DE FIGURAS ──────────────────────────────────────────────────
    (85,  'Arquitectura tecnológica propuesta - SIGA Odontología',
          'Arquitectura tecnológica propuesta - Sistema Gestión de Tareas'),
    (86,  'Arquitectura tecnológica actual - SIGA Odontología',
          'Arquitectura tecnológica actual - Sistema Gestión de Tareas EPO'),
    (112, 'Sistema SIGA Odontología me- diante metodología',
          'Sistema Gestión de Tareas EPO mediante metodología'),
    (112, 'Sistema SIGA Odontología mediante metodología',
          'Sistema Gestión de Tareas EPO mediante metodología'),

    # ── CASOS DE PRUEBA — RHF + Zod ────────────────────────────────────────
    (1093, 'El sistema de validación (Zod + React Hook Form) debe estar configu- rado.',
           'El sistema de autenticación JWT debe estar activo y configurado.'),
    (1093, 'El sistema de validación (Zod + React Hook Form) debe estar configurado.',
           'El sistema de autenticación JWT debe estar activo y configurado.'),
    (1218, 'El sistema de validación Zod + React Hook Form debe es- tar configu- rado.',
           'El sistema de autenticación JWT debe estar activo y configurado.'),
    (1218, 'El sistema de validación Zod + React Hook Form debe estar configurado.',
           'El sistema de autenticación JWT debe estar activo y configurado.'),
    (1256, 'El sistema renderiza formulario completo con React Hook Form.',
           'El sistema renderiza el formulario de nueva tarea correctamente.'),

    # ── CASOS DE PRUEBA — Multer ────────────────────────────────────────────
    (1844, 'El sistema de carga de imágenes (multer o MongoDB Atlas) de- be estar configu- rada.',
           'El sistema debe tener al menos una tarea registrada en la base de datos.'),
    (1844, 'El sistema de carga de imágenes (multer o MongoDB Atlas) debe estar configurada.',
           'El sistema debe tener al menos una tarea registrada en la base de datos.'),
    (1844, 'El sistema de carga de imágenes (multer o MongoDB Atlas) de- be estar configura',
           'El sistema debe tener al menos una tarea registrada.'),
    (1948, 'El backend procesa imagen con multer.',
           'El backend procesa la solicitud y registra la tarea correctamente.'),
    (2117, 'El sistema de carga de imágenes debe estar configurado (multer/MongoDB Atlas).',
           'El sistema debe tener al menos una tarea registrada en la base de datos.'),
    (2117, 'El sistema de carga de imágenes debe estar configurado (multer/MongoDB At- las).',
           'El sistema debe tener al menos una tarea registrada en la base de datos.'),

    # ── ÁRBOL DE ARCHIVOS ──────────────────────────────────────────────────
    (3527, 'multer.middleware.js  File Upload',
           'auth.middleware.js  JWT Authentication'),
    (3527, 'multer.middleware.js',
           'auth.middleware.js'),
    (3699, 'nodemailer  Email Client',
           'date-fns  Date Utilities'),
    (3699, 'nodemailer',
           'date-fns'),

    # ── HARDWARE — Nginx ────────────────────────────────────────────────────
    (3945, '2 vCPU dedicados (o equivalentes) para Node.js y Nginx.',
           '2 vCPU equivalentes gestionados por Render (plan gratuito cloud).'),
    (3945, 'Node.js y Nginx',
           'Node.js en Render'),

    # ── TABLA SOFTWARE — RHF+Zod, OpenAPI ─────────────────────────────────
    (4019, 'RHF + Zod',         'jsPDF-autotable'),
    (4019, 'React Hook Form',   'jsPDF'),
    (4036, 'Git + GitHub (proxy)',
           'Git + GitHub'),
    (4036, 'Proxy reverso y estáticos en producción.',
           'Control de versiones y repositorio del código fuente.'),
    (4039, 'OpenAPI 3 (opcio- nal)',  'Lucide React'),
    (4039, 'OpenAPI 3 (opcional)',    'Lucide React'),
    (4039, 'OpenAPI 3',              'Lucide React'),
    (4040, 'Especificación de endpoints y contratos de la API.',
           'Biblioteca de iconos SVG listos para React.'),
    (4040, 'Especificación de endpoints',
           'Biblioteca de iconos SVG'),
]

print("=== CORRECCIONES XML DIRECTAS ===")
for idx, old, new in fixes:
    if idx < len(paras):
        ok = replace_in_para_xml(paras[idx], old, new)
        status = '✓' if ok else '✗'
        print(f"  {status} [{idx}] '{old[:50]}' -> '{new[:40]}'")

# ── Buscar globalmente captions de figuras con "SIGA" ─────────────────────
print("\n--- Figuras/captions SIGA ---")
for i, para in enumerate(paras):
    full = para_full_text(para)
    if 'SIGA Odontolog' in full:
        ok = replace_in_para_xml(para, 'SIGA Odontología', 'Sistema Gestión de Tareas EPO')
        ok2 = replace_in_para_xml(para, 'SIGA Odontolog\u00eda', 'Sistema Gestión de Tareas EPO')
        if ok or ok2:
            print(f"  [{i}] -> {para_full_text(para)[:70]}")

# ── Buscar globalmente "File Upload" después del cambio de multer ─────────
print("\n--- File Upload residual ---")
for i, para in enumerate(paras):
    if 'File Upload' in para_full_text(para):
        ok = replace_in_para_xml(para, 'File Upload', 'JWT Authentication')
        if ok:
            print(f"  [{i}] -> {para_full_text(para)[:60]}")

# ── Buscar "Email Client" residual ────────────────────────────────────────
for i, para in enumerate(paras):
    if 'Email Client' in para_full_text(para):
        ok = replace_in_para_xml(para, 'Email Client', 'Date Utilities')
        if ok:
            print(f"  [{i}] -> {para_full_text(para)[:60]}")

# ── También en tablas ──────────────────────────────────────────────────────
print("\n--- Tablas residuales ---")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = para_full_text(para)
                changed = False
                for old, new in [
                    ('Servidor VPS básico (4 meses)', 'Render (backend cloud, plan gratuito)'),
                    ('Servidor VPS basico (4 meses)', 'Render (backend cloud, plan gratuito)'),
                    ('240.00', '0.00'),
                    ('Multer', 'jsPDF'),
                    ('multer', 'jsPDF'),
                    ('RHF + Zod', 'jsPDF-autotable'),
                    ('React Hook Form', 'jsPDF-autotable'),
                    ('Nginx', 'Git+GitHub'),
                    ('OpenAPI', 'Lucide React'),
                    ('Cloudinary', 'Recharts'),
                    ('Nodemailer', 'date-fns'),
                    ('Docker', 'Render'),
                    ('16 semanas', '13 semanas'),
                ]:
                    if old in full:
                        ok = replace_in_para_xml(para, old, new)
                        if ok:
                            changed = True
                            print(f"  Tabla: '{old[:30]}' -> '{new[:25]}'")
                            full = para_full_text(para)

doc.save(DEST)
print("\n✅ v3c guardado.")

# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA.docx")

# 1. Corregir estilo de los párrafos restaurados (Normal -> Body Text)
for p in doc.paragraphs:
    if ("Diagrama de Casos de Uso describe las interacciones" in p.text or
        "Diagramas de Secuencia UML documentan el flujo temporal" in p.text):
        p.style = doc.styles["Body Text"]

# 2. Eliminar párrafos vacíos antes de CAPÍTULO I
to_delete = []
for i, p in enumerate(doc.paragraphs):
    if "CAPÍTULO I" in p.text or "CAPITULO I" in p.text.upper():
        j = i - 1
        while j >= 0 and doc.paragraphs[j].text.strip() == "":
            to_delete.append(doc.paragraphs[j])
            j -= 1
        break

for p in to_delete:
    p._element.getparent().remove(p._element)

doc.save("INFORME_FINAL_SISTEMA_v2.docx")
print("Guardado como INFORME_FINAL_SISTEMA_v2.docx")

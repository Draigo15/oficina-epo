# -*- coding: utf-8 -*-
from docx import Document
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document("INFORME_FINAL_SISTEMA.docx")

reemplazos = [
    ("CS-08", "Gestionar Leads",         "CS-02", "Gestionar Tareas"),
    ("CS-09", "Gestionar Contenido",     "CS-03", "Generar Reporte PDF"),
    ("CS-04", "Enviar Mensaje de Con-",  "CS-04", "Gestionar Notificaciones"),
    ("CS-04", "Mnesaje de Contacto",     "CS-04", "Gestionar Notificaciones"),
    ("CS-04", "Enviar Mensaje de Contacto", "CS-04", "Gestionar Notificaciones"),
    ("Caso de Uso CS-08", "Gestionar Leads",    "Caso de Uso CS-02", "Gestionar Tareas"),
    ("Caso de Uso CS-09", "Gestionar Contenido","Caso de Uso CS-03", "Generar Reporte PDF"),
]

cambiados = 0
for p in doc.paragraphs:
    all_t = p._p.findall(f".//{{{NS}}}t")
    for t_el in all_t:
        if t_el.text:
            for (cs_old, desc_old, cs_new, desc_new) in reemplazos:
                if cs_old in t_el.text and desc_old in t_el.text:
                    t_el.text = t_el.text.replace(cs_old, cs_new).replace(desc_old, desc_new)
                    cambiados += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                all_t = p._p.findall(f".//{{{NS}}}t")
                for t_el in all_t:
                    if t_el.text:
                        for (cs_old, desc_old, cs_new, desc_new) in reemplazos:
                            if cs_old in t_el.text and desc_old in t_el.text:
                                t_el.text = t_el.text.replace(cs_old, cs_new).replace(desc_old, desc_new)
                                cambiados += 1

doc.save("INFORME_FINAL_SISTEMA.docx")
print(f"Renombrados {cambiados} elementos.")

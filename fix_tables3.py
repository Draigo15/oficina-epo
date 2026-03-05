# -*- coding: utf-8 -*-
from docx import Document

doc = Document("INFORME_Limpio.docx")

bad_terms = ["noticia", "institucional", "estudiante", "contacto", "programa", "progra-", "noti-", "informa-", "prospecto", "autenticar usua-", "conte-"]

eliminados_p = 0
for p in doc.paragraphs:
    text = p.text.lower()
    if p.text.strip().startswith("Narrativa") or p.text.strip().startswith("Caso de Uso") or p.text.strip().startswith("Diagrama de"):
        for bt in bad_terms:
            if bt in text:
                try:
                    p._element.getparent().remove(p._element)
                    eliminados_p += 1
                    break
                except:
                    pass

# Borrar filas sueltas de tablas que dicen esos terminos (como el indice o listas grandes)
for table in doc.tables:
    for row in table.rows:
        row_text = "".join(c.text.lower() for c in row.cells)
        if any(bt in row_text for bt in bad_terms) and ("cs-" in row_text or "cu-" in row_text):
            # No podemos borrar muy facil la fila limpia sin crashear XML en Docx a veces, intentemos borrar el texto
            for c in row.cells:
                c.text = ""

doc.save("INFORME_FINAL_SISTEMA.docx")
print("Hecho", eliminados_p)

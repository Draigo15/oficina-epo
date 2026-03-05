# -*- coding: utf-8 -*-
"""
Inserta narrativas reales del sistema TareasEpo antes de 'Requerimientos Funcionales'.
Casos de uso: CS-06, CS-01, CS-07, CS-02, CS-03, CS-04
"""
from docx import Document

doc = Document("INFORME_FINAL_SISTEMA_v2.docx")
paras = doc.paragraphs

# Encontrar parrafo ancla "Requerimientos Funcionales" (actualmente ~[649])
anchor_idx = None
for i, p in enumerate(paras):
    if p.text.strip() == "Requerimientos Funcionales" and p.style.name == "Heading 3":
        anchor_idx = i
        break

if anchor_idx is None:
    print("ERROR: No se encontro 'Requerimientos Funcionales'")
    exit(1)

print(f"Ancla encontrada en [{anchor_idx}]: '{paras[anchor_idx].text[:60]}'")
anchor_para = paras[anchor_idx]

# Helper: inserta parrafo con estilo ANTES del ancla
def ins(text, style):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    anchor_para._p.addprevious(p._p)
    return p

# Helper: parrafo en blanco
def blank():
    ins("", "Normal")

# ===================================================================
# ENCABEZADO GENERAL
# ===================================================================
ins("Narrativas de Casos de Uso", "Heading 3")
ins("A continuacion se presentan las narrativas detalladas de los casos de uso del "
    "Sistema de Gestion de Tareas - Oficina EPO, describiendo el flujo de interaccion "
    "entre los actores y el sistema para cada funcionalidad principal.", "Body Text")

# ===================================================================
# CS-06: AUTENTICAR USUARIO
# ===================================================================
ins("Narrativa del Caso de Uso CS-06 - Autenticar Usuario", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-06 | Nombre: Autenticar Usuario | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC / Asistente Administrativo", "Body Text")
ins("Precondicion: El usuario posee credenciales validas registradas en la coleccion "
    "Users de MongoDB Atlas. El servidor backend esta activo en Render.", "Body Text")
ins("Postcondicion: JWT generado y almacenado en localStorage. Usuario redirigido al Dashboard.", "Body Text")
ins("Disparador: El usuario accede a la URL del sistema (TareasEpo en Vercel) por primera vez "
    "o tras cerrar sesion.", "Body Text")
ins("Frecuencia: Alta - multiples veces al dia por cada usuario del sistema.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. El usuario accede a la URL del sistema desplegado en Vercel.", "Body Text")
ins("El sistema renderiza el componente Login.jsx con formulario de email y contrasena.", "List Paragraph")
ins("2. El usuario ingresa su email y contrasena en el formulario.", "Body Text")
ins("El sistema habilita el boton 'Iniciar Sesion' al detectar datos en ambos campos.", "List Paragraph")
ins("3. El usuario hace clic en el boton 'Iniciar Sesion'.", "Body Text")
ins("El sistema ejecuta peticion POST /api/auth/login con payload {email, password}.", "List Paragraph")
ins("El sistema aplica bcryptjs.compare() para verificar el hash de la contrasena.", "List Paragraph")
ins("El sistema verifica que el campo activo del usuario sea true.", "List Paragraph")
ins("El sistema genera un JWT firmado con la clave secreta, con payload {id, nombre, rol} "
    "y expiracion de 7 dias.", "List Paragraph")
ins("4. El sistema retorna respuesta 200 con el JWT y datos del usuario.", "Body Text")
ins("El frontend almacena el JWT en localStorage con la clave 'user'.", "List Paragraph")
ins("AuthContext.jsx actualiza el estado global con los datos del usuario autenticado.", "List Paragraph")
ins("PrivateRoute.jsx habilita el acceso a rutas protegidas.", "List Paragraph")
ins("5. El sistema redirige al usuario al Dashboard (/dashboard).", "Body Text")
ins("React Router DOM renderiza el componente Dashboard.jsx segun el rol del usuario.", "List Paragraph")
ins("El sistema muestra las secciones disponibles segun el rol: "
    "Jefa (acceso completo) o Asistente (acceso restringido).", "List Paragraph")
blank()

ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Ruta backend: POST /api/auth/login (routes/auth.js)", "Body Text")
ins("Middleware de autenticacion: middleware/auth.js - verifica JWT en header Authorization", "Body Text")
ins("Modelo de datos: models/User.js - coleccion 'users' en MongoDB Atlas", "Body Text")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("Token JWT valido almacenado en localStorage del navegador del usuario.", "Body Text")
ins("El estado de AuthContext contiene: {id, nombre, email, rol, token}.", "Body Text")
ins("Todas las rutas protegidas de la aplicacion son accesibles para el usuario.", "Body Text")
ins("Se registra timestamp de ultimo acceso en el documento del usuario en MongoDB.", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Credenciales incorrectas - el sistema retorna error 401 y muestra mensaje "
    "'Email o contrasena incorrectos' mediante ToastContext.", "Body Text")
ins("FA-02: Usuario inactivo - el campo activo=false en MongoDB; el sistema retorna "
    "error 403 y muestra 'Cuenta desactivada. Contacte al administrador'.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Error de conectividad con backend (Render inactivo/timeout) - el sistema "
    "muestra 'Error de conexion. Intente nuevamente' con boton de reintento.", "Body Text")
ins("FE-02: JWT expirado durante sesion activa - AuthContext detecta error 401 en "
    "llamada api.js, limpia localStorage y redirige a /login automaticamente.", "Body Text")
blank()

# ===================================================================
# CS-01: VISUALIZAR DASHBOARD
# ===================================================================
ins("Narrativa del Caso de Uso CS-01 - Visualizar Dashboard", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-01 | Nombre: Visualizar Dashboard | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC / Asistente Administrativo", "Body Text")
ins("Precondicion: Usuario autenticado con JWT valido. Token presente en localStorage.", "Body Text")
ins("Postcondicion: Panel principal visible con tarjetas de resumen cargadas desde la API.", "Body Text")
ins("Disparador: Acceso a la ruta /dashboard tras autenticacion exitosa.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. El usuario accede al Dashboard tras iniciar sesion.", "Body Text")
ins("Layout.jsx renderiza la barra lateral de navegacion y el area de contenido principal.", "List Paragraph")
ins("El sistema ejecuta peticion GET /api/tasks/stats con JWT en header Authorization.", "List Paragraph")
ins("2. El sistema obtiene las estadisticas del area desde MongoDB.", "Body Text")
ins("El sistema calcula: total de tareas, tareas pendientes, tareas en progreso, "
    "tareas completadas, tareas de alta prioridad.", "List Paragraph")
ins("El sistema agrupa tareas por semana para grafico de barras con Recharts.", "List Paragraph")
ins("El sistema calcula distribucion de tareas por prioridad para grafico de dona.", "List Paragraph")
ins("3. El sistema renderiza el Dashboard con los datos obtenidos.", "Body Text")
ins("Se muestran 5 tarjetas de resumen (KPI cards) con contadores en tiempo real.", "List Paragraph")
ins("Se renderiza grafico de barras (BarChart) mostrando tareas completadas por semana.", "List Paragraph")
ins("Se renderiza grafico de dona (PieChart) con distribucion por prioridad.", "List Paragraph")
ins("4. El usuario visualiza las notificaciones recientes no leidas.", "Body Text")
ins("El sistema ejecuta GET /api/notifications?limit=5 y muestra las ultimas 5.", "List Paragraph")
ins("Se muestra badge con conteo de notificaciones no leidas en el menu lateral.", "List Paragraph")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("El dashboard muestra estadisticas actualizadas reflejo del estado real de MongoDB.", "Body Text")
ins("Los graficos Recharts son interactivos: tooltip al hover sobre barras y sectores.", "Body Text")
ins("El usuario puede navegar a cualquier modulo desde el menu lateral de Layout.jsx.", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Sin tareas registradas - el sistema muestra valores cero en KPI cards y "
    "graficos vacios con mensaje 'No hay tareas registradas aun'.", "Body Text")
ins("FA-02: Acceso con rol Asistente - el sistema oculta las opciones de gestion de "
    "usuarios y configuracion avanzada en el menu lateral.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Error al cargar estadisticas - el sistema muestra skeleton loaders durante "
    "la carga y Toast de error si la peticion falla tras 3 reintentos.", "Body Text")
blank()

# ===================================================================
# CS-02: GESTIONAR TAREAS
# ===================================================================
ins("Narrativa del Caso de Uso CS-02 - Gestionar Tareas", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-02 | Nombre: Gestionar Tareas | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC. Actor Secundario: Asistente Administrativo.", "Body Text")
ins("Precondicion: Usuario autenticado. La coleccion 'tasks' existe en MongoDB Atlas.", "Body Text")
ins("Postcondicion: Tarea creada, editada, eliminada o cambiada de estado en MongoDB.", "Body Text")
ins("Disparador: El usuario accede al modulo /tasks desde el menu lateral.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. El usuario accede al modulo de Tareas.", "Body Text")
ins("El sistema ejecuta GET /api/tasks con JWT. Recupera todas las tareas de la coleccion.", "List Paragraph")
ins("Tasks.jsx renderiza el panel Kanban con tres columnas: Pendiente, En Progreso, Completada.", "List Paragraph")
ins("Las tareas se muestran como tarjetas con titulo, descripcion, prioridad (color), "
    "responsable y fecha limite.", "List Paragraph")
ins("2. La Jefa hace clic en 'Nueva Tarea'.", "Body Text")
ins("El sistema abre modal de creacion con campos: titulo (requerido), descripcion, "
    "prioridad (alta/media/baja), responsable (lista de usuarios), fecha limite.", "List Paragraph")
ins("3. La Jefa completa el formulario y hace clic en 'Guardar'.", "Body Text")
ins("El sistema ejecuta POST /api/tasks con los datos del formulario.", "List Paragraph")
ins("MongoDB inserta el nuevo documento en la coleccion 'tasks' con timestamps.", "List Paragraph")
ins("El sistema envia notificacion automatica al responsable asignado via POST /api/notifications.", "List Paragraph")
ins("El panel Kanban se actualiza en tiempo real mostrando la nueva tarea en columna 'Pendiente'.", "List Paragraph")
ins("4. El usuario arrastra una tarea a otra columna para cambiar su estado.", "Body Text")
ins("React Beautiful DnD detecta el evento onDragEnd.", "List Paragraph")
ins("El sistema ejecuta PUT /api/tasks/:id con el nuevo estado.", "List Paragraph")
ins("MongoDB actualiza el campo 'estado' y 'updatedAt' del documento.", "List Paragraph")
ins("5. La Jefa elimina una tarea haciendo clic en el icono de eliminar.", "Body Text")
ins("El sistema muestra modal de confirmacion: 'Esta seguro de eliminar esta tarea?'.", "List Paragraph")
ins("Confirmado, el sistema ejecuta DELETE /api/tasks/:id.", "List Paragraph")
ins("MongoDB elimina el documento. El panel Kanban se actualiza removiendo la tarjeta.", "List Paragraph")
ins("El sistema muestra Toast de confirmacion: 'Tarea eliminada exitosamente'.", "List Paragraph")
blank()

ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Ruta backend: GET|POST /api/tasks - PUT|DELETE /api/tasks/:id (routes/tasks.js)", "Body Text")
ins("Modelo de datos: models/Task.js - coleccion 'tasks' con campos: titulo, descripcion, "
    "estado, prioridad, responsable, fechaLimite, creadoPor, timestamps.", "Body Text")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("Los cambios quedan persistidos en MongoDB Atlas (coleccion 'tasks').", "Body Text")
ins("El historial de cambios de estado es trazable mediante timestamps automaticos.", "Body Text")
ins("Las estadisticas del Dashboard se actualizan al recargar.", "Body Text")
ins("El responsable asignado recibe notificacion interna en su bandeja.", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Filtrar tareas - el usuario aplica filtros por estado, prioridad o responsable; "
    "el sistema actualiza la vista Kanban sin recargar pagina.", "Body Text")
ins("FA-02: Busqueda textual - el usuario escribe en la barra de busqueda; "
    "el sistema filtra tarjetas en tiempo real por coincidencia en titulo o descripcion.", "Body Text")
ins("FA-03: Editar tarea - el usuario hace clic en 'Editar'; el sistema abre el modal "
    "precargado con los datos actuales y ejecuta PUT /api/tasks/:id al guardar.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Token expirado durante operacion - la API retorna 401; el sistema limpia "
    "sesion y redirige a /login con mensaje 'Sesion expirada'.", "Body Text")
ins("FE-02: Error de validacion - titulo vacio o fecha limite invalida; "
    "el sistema muestra errores inline en el formulario sin cerrar el modal.", "Body Text")
blank()

# ===================================================================
# CS-03: GENERAR REPORTE PDF
# ===================================================================
ins("Narrativa del Caso de Uso CS-03 - Generar Reporte PDF", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-03 | Nombre: Generar Reporte PDF | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC.", "Body Text")
ins("Precondicion: Usuario autenticado con rol Jefa. Tareas registradas en el periodo seleccionado.", "Body Text")
ins("Postcondicion: Archivo PDF generado y descargado en el dispositivo del usuario.", "Body Text")
ins("Disparador: La Jefa accede al modulo /reports y selecciona un periodo mensual.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. La Jefa accede al modulo de Reportes desde el menu lateral.", "Body Text")
ins("Reports.jsx renderiza el selector de periodo (mes/anio picker) y opciones de formato.", "List Paragraph")
ins("2. La Jefa selecciona el mes y anio del reporte y hace clic en 'Generar Reporte'.", "Body Text")
ins("El sistema ejecuta GET /api/reports/monthly?mes=MM&anio=YYYY con JWT.", "List Paragraph")
ins("El backend consulta MongoDB filtrando documentos de la coleccion 'tasks' "
    "donde fechaCreacion este en el rango del periodo.", "List Paragraph")
ins("El backend calcula estadisticas: total tareas, completadas, pendientes, en progreso, "
    "tasa de completitud y tiempo promedio de resolucion.", "List Paragraph")
ins("3. El sistema retorna los datos del reporte en formato JSON.", "Body Text")
ins("Reports.jsx recibe los datos y activa la generacion del PDF en el cliente.", "List Paragraph")
ins("jsPDF instancia el documento PDF con orientacion vertical y formato A4.", "List Paragraph")
ins("jsPDF-autotable genera la tabla de tareas con columnas: N, Titulo, Responsable, "
    "Estado, Prioridad, Fecha Limite, Fecha Completada.", "List Paragraph")
ins("Se incluye encabezado con: logo institucional (si disponible), titulo del reporte, "
    "periodo seleccionado, fecha de generacion y nombre del usuario.", "List Paragraph")
ins("Se incluye resumen estadistico en la seccion superior: KPIs del periodo.", "List Paragraph")
ins("4. El PDF se descarga automaticamente en el navegador.", "Body Text")
ins("jsPDF ejecuta doc.save('Reporte_CMC_MM_YYYY.pdf') activando la descarga.", "List Paragraph")
ins("El sistema muestra Toast: 'Reporte generado exitosamente'.", "List Paragraph")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("Archivo PDF descargado localmente con nombre 'Reporte_CMC_[mes]_[anio].pdf'.", "Body Text")
ins("El reporte refleja fielmente los datos de MongoDB del periodo seleccionado.", "Body Text")
ins("La generacion es completamente del lado cliente (sin carga adicional al servidor).", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Sin tareas en el periodo - el sistema muestra mensaje 'No hay tareas "
    "registradas para el periodo seleccionado' y deshabilita el boton de generar.", "Body Text")
ins("FA-02: Vista previa - el usuario selecciona 'Vista Previa'; el sistema muestra "
    "el PDF embebido en un iframe antes de la descarga.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Error al obtener datos del backend - el sistema muestra Toast de error "
    "y permite reintentar sin perder la seleccion del periodo.", "Body Text")
blank()

# ===================================================================
# CS-04: GESTIONAR NOTIFICACIONES
# ===================================================================
ins("Narrativa del Caso de Uso CS-04 - Gestionar Notificaciones", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-04 | Nombre: Gestionar Notificaciones | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC. Actor Secundario: Asistente Administrativo (receptor).", "Body Text")
ins("Precondicion: Usuarios autenticados. Coleccion 'notifications' en MongoDB Atlas.", "Body Text")
ins("Postcondicion: Notificacion creada, leida o eliminada. Contadores actualizados.", "Body Text")
ins("Disparador: Asignacion de tarea (automatica) o acceso al modulo /notifications.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. El usuario accede al Centro de Notificaciones desde el icono en la barra superior.", "Body Text")
ins("NotificationsCenter.jsx ejecuta GET /api/notifications con JWT.", "List Paragraph")
ins("El sistema retorna todas las notificaciones del usuario autenticado, ordenadas por fecha.", "List Paragraph")
ins("Se muestra badge rojo con el conteo de notificaciones no leidas.", "List Paragraph")
ins("2. El sistema renderiza la lista de notificaciones.", "Body Text")
ins("Cada notificacion muestra: tipo (nueva tarea, recordatorio, sistema), mensaje, "
    "fecha y estado (leida/no leida).", "List Paragraph")
ins("Las notificaciones no leidas se destacan con fondo diferenciado.", "List Paragraph")
ins("3. El usuario hace clic en una notificacion no leida.", "Body Text")
ins("El sistema ejecuta PATCH /api/notifications/:id/read para marcarla como leida.", "List Paragraph")
ins("MongoDB actualiza el campo 'leida' a true y registra 'fechaLeida': timestamp.", "List Paragraph")
ins("El badge de notificaciones no leidas decrementa su conteo.", "List Paragraph")
ins("4. La Jefa crea una notificacion manual para el equipo.", "Body Text")
ins("La Jefa accede a 'Nueva Notificacion' y completa: destinatario, tipo, mensaje.", "List Paragraph")
ins("El sistema ejecuta POST /api/notifications con los datos.", "List Paragraph")
ins("MongoDB inserta el documento en la coleccion 'notifications'.", "List Paragraph")
ins("El destinatario ve el badge incrementado en su proxima carga.", "List Paragraph")
blank()

ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Ruta backend: GET /api/notifications - POST /api/notifications - "
    "PATCH /api/notifications/:id/read (routes/notifications.js)", "Body Text")
ins("Modelo de datos: models/Notification.js - coleccion 'notifications' con campos: "
    "destinatario (ObjectId ref User), tipo, mensaje, leida (Boolean), fechaCreacion.", "Body Text")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("Los cambios de estado (leida/no leida) quedan persistidos en MongoDB.", "Body Text")
ins("Los contadores del badge se mantienen sincronizados con el estado real.", "Body Text")
ins("La Jefa tiene trazabilidad de cuando cada notificacion fue leida.", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Marcar todas como leidas - el usuario hace clic en 'Marcar todas como leidas'; "
    "el sistema ejecuta PATCH /api/notifications/readAll y actualiza todos los documentos.", "Body Text")
ins("FA-02: Eliminar notificacion - el usuario hace clic en el icono eliminar; "
    "el sistema ejecuta DELETE /api/notifications/:id y remueve la notificacion de la lista.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Sin notificaciones - el sistema muestra ilustracion y mensaje "
    "'No tienes notificaciones pendientes'.", "Body Text")
blank()

# ===================================================================
# CS-07: ADMINISTRAR DASHBOARD (Estadisticas avanzadas)
# ===================================================================
ins("Narrativa del Caso de Uso CS-07 - Administrar Dashboard", "Heading 3")
ins("Campo/Seccion        Descripcion/Contenido", "Heading 3")
ins("Caso de Uso: CS-07 | Nombre: Administrar Dashboard | Version: 1.0", "Body Text")
ins("Actor Principal: Jefa del CMC.", "Body Text")
ins("Precondicion: Usuario autenticado con rol Jefa. Tareas registradas en la base de datos.", "Body Text")
ins("Postcondicion: Estadisticas avanzadas visualizadas y usuario actualiza perfil o usuarios.", "Body Text")
ins("Disparador: La Jefa accede al modulo de estadisticas /stats desde el menu lateral.", "Body Text")
blank()

ins("NARRATIVA PRINCIPAL", "Heading 2")
ins("Accion del Actor        Respuesta del Sistema", "Heading 3")
ins("1. La Jefa accede al modulo de Estadisticas.", "Body Text")
ins("Stats.jsx ejecuta multiples peticiones en paralelo mediante Axios:", "List Paragraph")
ins("GET /api/tasks/stats - estadisticas generales de tareas.", "List Paragraph")
ins("GET /api/tasks/byUser - tareas agrupadas por responsable.", "List Paragraph")
ins("GET /api/tasks/trend - evolucion temporal de carga de trabajo.", "List Paragraph")
ins("2. El sistema renderiza el dashboard de estadisticas avanzadas.", "Body Text")
ins("Grafico de lineas (LineChart de Recharts): evolucion de tareas completadas vs. creadas por semana.", "List Paragraph")
ins("Grafico de barras apiladas (StackedBarChart): carga de trabajo por responsable.", "List Paragraph")
ins("Grafico de dona (PieChart): distribucion de tareas por estado y prioridad.", "List Paragraph")
ins("Tabla resumen con KPIs: tasa de completitud, tiempo promedio de resolucion, "
    "responsable con mayor productividad.", "List Paragraph")
ins("3. La Jefa gestiona cuentas de usuario desde el submenu de administracion.", "Body Text")
ins("El sistema muestra tabla de usuarios con columnas: nombre, email, rol, activo/inactivo.", "List Paragraph")
ins("La Jefa puede crear nuevo usuario (POST /api/auth/register) o "
    "desactivar cuenta existente (PUT /api/users/:id).", "List Paragraph")
blank()

ins("POSTCONDICIONES", "Heading 2")
ins("La Jefa tiene vision completa del rendimiento del area y puede tomar decisiones informadas.", "Body Text")
ins("Los cambios en cuentas de usuario quedan persistidos en MongoDB Atlas.", "Body Text")
blank()

ins("FLUJOS ALTERNATIVOS", "Heading 2")
ins("FA-01: Filtrar por rango de fechas - la Jefa selecciona un rango personalizado; "
    "el sistema recupera estadisticas del intervalo especificado.", "Body Text")
ins("FA-02: Exportar estadisticas - la Jefa genera reporte PDF con los graficos del periodo.", "Body Text")
blank()

ins("FLUJOS DE EXCEPCION", "Heading 2")
ins("FE-01: Intento de crear usuario con email duplicado - la API retorna 400; "
    "el sistema muestra 'El email ya esta registrado en el sistema'.", "Body Text")
blank()

print(f"Narrativas insertadas correctamente antes de [{anchor_idx}].")
print(f"Total parrafos ahora: {len(doc.paragraphs)}")

out = "INFORME_FINAL_SISTEMA_v2.docx"
try:
    doc.save(out)
    print(f"Guardado en {out}")
except PermissionError:
    out = "INFORME_FINAL_SISTEMA_v3.docx"
    doc.save(out)
    print(f"Bloqueado, guardado en {out}")

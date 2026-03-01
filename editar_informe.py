"""
Script para adaptar el INFORME.docx del sistema de Odontología
al Sistema de Gestión de Tareas - Oficina EPO
Autor adaptado: Lira Alvarez, Rodrigo Samael Adonai (2019063331)
"""

import shutil
import os
from docx import Document
from docx.shared import Pt
import re

# ─── Rutas ───────────────────────────────────────────────────────────────────
SRC  = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME.docx'
BAK  = r'c:\Users\carus\OneDrive\Escritorio\PRACTICAS\TareasEpo\INFORME_BACKUP.docx'
DEST = SRC   # sobreescribir el mismo archivo

# Crear backup si no existe
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print(f"✓ Backup creado: INFORME_BACKUP.docx")
else:
    print(f"  Backup ya existe, se omite.")

doc = Document(SRC)

# ─── Utilidades ──────────────────────────────────────────────────────────────

def replace_text_in_para(para, old, new):
    """Reemplaza texto en un párrafo preservando el estilo del primer run."""
    full = para.text
    if old.lower() in full.lower():
        # Limpiar todos los runs y poner el nuevo texto en el primero
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = full.replace(old, new)
            else:
                run.text = ''
        return True
    return False

def replace_in_doc(old, new, max_replace=999):
    """Reemplaza todas las ocurrencias de 'old' por 'new' en el documento."""
    count = 0
    for para in doc.paragraphs:
        if old in para.text and count < max_replace:
            # Reemplazar en cada run preservando formato
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
    return count

def set_para_text(para, new_text):
    """Establece el texto completo de un párrafo (limpia runs extras)."""
    for i, run in enumerate(para.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ''
    if not para.runs:
        para.add_run(new_text)

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 1 — PORTADA: Título del sistema
# ═══════════════════════════════════════════════════════════════════════════════
old_titulo  = 'Desarrollo de un Sistema Web de Gestión Académica para la Escuela Profesional de Odontología'
new_titulo  = 'Desarrollo de un Sistema Web de Gestión de Tareas para la Oficina del Comité de Mejora Continua de la Escuela Profesional de Odontología'

old_titulo2 = '\u201cDesarrollo de un Sistema Web de Gestión Académica para la Escuela Profesional de Odontología\u201d'
new_titulo2 = '\u201cDesarrollo de un Sistema Web de Gestión de Tareas para la Oficina del Comité de Mejora Continua de la Escuela Profesional de Odontología\u201d'

changed = 0
for para in doc.paragraphs:
    t = para.text
    if 'Gestión Académica para la Escuela Profesional de Odontología' in t and 'Desarrollo de un Sistema' in t:
        set_para_text(para, new_titulo2 if '"' in t or '\u201c' in t else new_titulo)
        changed += 1
        break

print(f"  Portada título: {changed} cambio(s)")

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 2 — Nombre del sistema en todo el documento
# ═══════════════════════════════════════════════════════════════════════════════
reemplazos_globales = [
    # Sistema completo
    ("SIGA Odontología", "Sistema de Gestión de Tareas – Oficina EPO"),
    ("SIGA Odontologia", "Sistema de Gestión de Tareas – Oficina EPO"),
    ("Sistema Integral de Gestión Académica (SIGA Odontología)",
     "Sistema de Gestión de Tareas – Oficina EPO"),
    ("Sistema Integral de Gestión Académica (SIGA Odontologia)",
     "Sistema de Gestión de Tareas – Oficina EPO"),
    # Autor amigo
    ("Vela Vargas, Abraham Jesús", "Lira Alvarez, Rodrigo Samael Adonai"),
    ("Abraham Jesús Vela Vargas",  "Rodrigo Samael Adonai Lira Alvarez"),
    ("2019063322", "2019063331"),
]

for old, new in reemplazos_globales:
    n = replace_in_doc(old, new)
    if n:
        print(f"  Reemplazado '{old[:50]}': {n} vez/veces")

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 3 — INTRODUCCIÓN (párrafos 155-161 aprox.)
# Buscar por texto inicial único y reemplazar el bloque
# ═══════════════════════════════════════════════════════════════════════════════
intro_nuevo = [
    "El presente informe de prácticas pre-profesionales de la carrera de Ingeniería de Sistemas tiene como objetivo documentar el desarrollo y la aplicación de conocimientos técnicos especializados adquiridos durante la formación académica, materializados en la implementación de una solución tecnológica para el área administrativa de la Escuela Profesional de Odontología. Las actividades se desarrollaron enfocadas en el diseño, desarrollo e implementación de un sistema web orientado a automatizar la gestión de tareas y la generación de reportes mensuales en el Comité de Mejora Continua (CMC) de la institución.",

    "Durante esta experiencia profesional, se ejecutó el análisis, diseño, desarrollo e implementación de una aplicación web full-stack denominada Sistema de Gestión de Tareas – Oficina EPO, empleando el stack tecnológico MERN (MongoDB, Express.js, React y Node.js). El sistema abarca módulos especializados para la creación y seguimiento de tareas con prioridades diferenciadas, registro automático de fecha y hora de completado, generación automática de reportes mensuales en formato PDF, un panel de control con estadísticas en tiempo real, y un sistema de notificaciones interno, todo ello bajo un esquema de autenticación y autorización diferenciada por roles (Jefa y Asistente).",

    "A través del proceso de levantamiento de requerimientos y análisis de necesidades del área, se identificó la ausencia de una plataforma digital que permitiera centralizar y registrar las tareas asignadas al personal de la oficina, controlar su estado de cumplimiento y generar evidencia documental mensual de las actividades realizadas. Esta problemática motivó el desarrollo de una arquitectura cliente-servidor moderna con API RESTful, implementando buenas prácticas de seguridad mediante encriptación de contraseñas con bcrypt, autenticación con JSON Web Tokens (JWT), y una interfaz de usuario responsive desarrollada con Tailwind CSS y React 18 que garantiza accesibilidad desde cualquier dispositivo.",

    "El presente documento detalla el proceso completo de ingeniería de software aplicado, desde el modelado de base de datos NoSQL con tres colecciones especializadas (usuarios, tareas, notificaciones) hasta la implementación de los endpoints de la API, destacando la aplicación práctica de metodologías de desarrollo web modernas, arquitectura de software, gestión de estado con Context API, manejo de autenticación y autorización por roles, e integración de funcionalidades avanzadas como la generación de reportes PDF con jsPDF y jsPDF-autotable, visualización de estadísticas con Recharts, y organización de tareas con arrastrar y soltar mediante React Beautiful DnD. Este informe evidencia el impacto de la formación académica en ingeniería de sistemas para la generación de soluciones tecnológicas que digitalizan y optimizan procesos administrativos en el sector educativo.",
]

marcadores_intro = [
    "El presente informe de prácticas pre-profesionales de la carrera de Ingeniería",
    "presente informe de pr",  # forma corta por si tiene encoding raro
]

intro_idx = -1
for i, para in enumerate(doc.paragraphs):
    if any(m.lower()[:30] in para.text.lower() for m in marcadores_intro):
        # Verificar que esté en la sección de introducción (no otro lugar)
        if i > 150 and i < 200:
            intro_idx = i
            break

if intro_idx >= 0:
    # Los párrafos de intro son 155,156,157,158 aprox. (4 párrafos de texto)
    # Reemplazar los 4 primeros párrafos de texto continuo
    idx = intro_idx
    nuevo_idx = 0
    while nuevo_idx < len(intro_nuevo) and idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        if para.text.strip():  # solo párrafos con texto
            set_para_text(para, intro_nuevo[nuevo_idx])
            nuevo_idx += 1
        idx += 1
    print(f"  Introducción: {nuevo_idx} párrafo(s) reemplazado(s) desde índice {intro_idx}")
else:
    print("  ADVERTENCIA: No se encontró la Introducción")

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 4 — CAPÍTULO III: ACTIVIDADES DESARROLLADAS
# ═══════════════════════════════════════════════════════════════════════════════

# Párrafo de apertura del Capítulo III (índice ~272)
cap3_apertura_nuevo = (
    "Durante mi experiencia laboral en la Escuela Profesional de Odontología de "
    "la Universidad Privada de Tacna, desarrollé funciones de soporte técnico en "
    "el área del Comité de Mejora Continua (CMC), orientadas al desarrollo e "
    "implementación de soluciones tecnológicas. Las actividades realizadas "
    "incluyeron el relevamiento de requerimientos con el personal administrativo "
    "para identificar necesidades de digitalización, el diseño y programación del "
    "Sistema de Gestión de Tareas – Oficina EPO, el registro y seguimiento de "
    "tareas institucionales, así como el apoyo en el procesamiento estadístico de "
    "datos académicos y la generación de reportes para la toma de decisiones."
)

cap3_funciones_nuevo = (
    "Durante mi desempeño en la Escuela Profesional de Odontología de la "
    "Universidad Privada de Tacna, mis responsabilidades abarcaron el desarrollo "
    "web full-stack, el soporte técnico informático y el análisis de datos "
    "institucionales. Las funciones principales incluyeron:"
)

# 3.4.1 - primera labor
labor1_titulo  = "Desarrollo e Implementación del Sistema de Gestión de Tareas"
labor1_cuerpo1 = (
    "Realicé el análisis, diseño, desarrollo e implementación del Sistema de "
    "Gestión de Tareas – Oficina EPO, una aplicación web full-stack orientada a "
    "centralizar y automatizar el registro de tareas del personal administrativo "
    "del Comité de Mejora Continua de la Escuela Profesional de Odontología."
)
labor1_cuerpo2 = (
    "El proceso de desarrollo se ejecutó aplicando la metodología Extreme "
    "Programming (XP) con iteraciones cortas, utilizando el stack tecnológico "
    "MERN (MongoDB, Express.js, React y Node.js), Tailwind CSS para la interfaz "
    "y JWT para autenticación segura diferenciada por roles."
)
labor1_pasos = (
    "Los procedimientos implementados incluyeron:\n"
    "Levantamiento de requerimientos: Reuniones con la Jefa de Oficina para "
    "definir funcionalidades, roles de usuario (Jefa y Asistente) y reglas de "
    "negocio del sistema.\n"
    "Diseño y desarrollo: Modelado de base de datos en MongoDB, implementación "
    "de API RESTful con Node.js/Express, y desarrollo de interfaz React con "
    "gestión de estado mediante Context API.\n"
    "Pruebas y despliegue: Validación funcional de módulos, corrección de errores "
    "y puesta en producción del backend en Render y frontend en Vercel."
)

# 3.4.2
labor2_titulo  = "Gestión y Seguimiento de Tareas Institucionales"
labor2_cuerpo1 = (
    "Ejecuté el registro, asignación y seguimiento de tareas del área mediante el "
    "sistema implementado, apoyando al personal en la adopción de la herramienta "
    "digital y verificando el correcto funcionamiento del flujo de trabajo."
)
labor2_pasos = (
    "Las actividades desarrolladas incluyeron:\n"
    "Registro de tareas: Alta de tareas con título, descripción, prioridad "
    "(normal o alta) y fecha límite, asegurando trazabilidad del trabajo diario.\n"
    "Seguimiento de estado: Monitoreo del estado de tareas (pendiente / "
    "completada) y verificación del registro automático de fecha y hora de "
    "completado.\n"
    "Soporte a usuarios: Capacitación básica al personal sobre el uso del sistema "
    "y resolución de incidencias operativas durante el periodo de adopción."
)

# 3.4.3
labor3_titulo  = "Procesamiento Estadístico y Generación de Reportes"
labor3_cuerpo1 = (
    "Realicé la configuración y validación del módulo de reportes mensuales "
    "del sistema, verificando la correcta generación automática de reportes en "
    "formato PDF con el listado de tareas completadas por período, así como el "
    "procesamiento de datos estadísticos institucionales en apoyo al CMC."
)
labor3_pasos = (
    "Los procedimientos implementados incluyeron:\n"
    "Validación de reportes PDF: Prueba del módulo de generación de reportes "
    "mensuales con jsPDF y jsPDF-autotable, verificando integridad de datos y "
    "formato de salida conforme a los requerimientos del área.\n"
    "Estadísticas del dashboard: Revisión de los indicadores del panel principal "
    "(tareas totales, pendientes, completadas, alta prioridad) con datos reales "
    "de prueba.\n"
    "Apoyo en análisis de datos: Colaboración en el procesamiento de información "
    "estadística de evaluaciones y asistencias del personal docente para "
    "generación de reportes administrativos del CMC."
)

# 3.4.4
labor4_titulo  = "Soporte Técnico Informático"
labor4_cuerpo1 = (
    "Brindé soporte técnico general al área del Comité de Mejora Continua, "
    "atendiendo requerimientos de mantenimiento de equipos, configuración de "
    "software y apoyo en la gestión de recursos digitales institucionales."
)
labor4_pasos = (
    "Las actividades desarrolladas incluyeron:\n"
    "Mantenimiento de equipos: Diagnóstico y resolución de fallas básicas de "
    "hardware y software en las computadoras del área.\n"
    "Configuración de software: Instalación y configuración de herramientas de "
    "productividad y entornos de desarrollo necesarios para las actividades "
    "del CMC.\n"
    "Gestión digital: Organización de archivos institucionales y apoyo en la "
    "actualización de registros digitales del área."
)

# ─── Buscar y reemplazar el bloque del Capítulo III ──────────────────────────
cap3_marcador = "CAPÍTULO III:"
cap3_marcador2 = "CAP\u00cdTULO III:"

cap3_idx = -1
for i, para in enumerate(doc.paragraphs):
    if "CAPÍTULO III" in para.text or "CAP" in para.text and "TULO III" in para.text:
        cap3_idx = i
        break

if cap3_idx >= 0:
    print(f"  Capítulo III encontrado en índice {cap3_idx}")

    # Mapeo de textos originales → nuevos (buscar por fragmento único)
    reemplazos_cap3 = {
        # Párrafo apertura ~272
        "Durante mi experiencia laboral en la Escuela Profesional de Odontología de la Universidad Privada de Tacna, desarrollé diversas funciones relacionadas con el diseño gráfico": cap3_apertura_nuevo,
        # Párrafo funciones ~279
        "Durante mi desempeño en la Escuela Profesional de Odontología de la Universidad Privada de Tacna, mis responsabilidades abarcaron diseño gráfico": cap3_funciones_nuevo,
        # Labores
        "Diseño Gráfico de Material Promocional Institucional": labor1_titulo,
        "Desarrollé material gráfico promocional para la difusión de acti- vidades académicas": labor1_cuerpo1,
        "Desarrollé material gráfico promocional para la difusión de actividades académicas": labor1_cuerpo1,
        "El proceso de diseño se ejecutó siguiendo lineamientos de iden- tidad visual institucional": labor1_cuerpo2,
        "El proceso de diseño se ejecutó siguiendo lineamientos de identidad visual institucional": labor1_cuerpo2,
        # Análisis de Datos de Asistencia Docente → Gestión y Seguimiento
        "Análisis de Datos de Asistencia Docente": labor2_titulo,
        "Ejecuté procesos de análisis de bases de datos institucionales para evaluar patrones de asistencia del personal docente": labor2_cuerpo1,
        # Procesamiento Estadístico
        "Procesamiento Estadístico de Evaluaciones Docentes": labor3_titulo,
        "Realicé recolección, tabulación y análisis estadístico de evaluaciones estudiantiles al desempeño docente": labor3_cuerpo1,
        "Realicé recolección, tabulación y anßlisis estadístico de evaluaciones estudiantiles": labor3_cuerpo1,
        # Gestión Web → Soporte Técnico
        "Gestión y Actualización de Contenido Web Institucional": labor4_titulo,
        "Administré y actualicé el contenido digital del sitio web oficial de la Escuela Profesional de Odontología": labor4_cuerpo1,
    }

    for i, para in enumerate(doc.paragraphs):
        if i < cap3_idx:
            continue
        txt = para.text
        for old_frag, new_text in reemplazos_cap3.items():
            if old_frag.lower()[:40] in txt.lower():
                set_para_text(para, new_text)
                print(f"    ✓ Reemplazado: '{old_frag[:50]}'")
                break
else:
    print("  ADVERTENCIA: No se encontró Capítulo III")

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 5 — CAPÍTULO IV: Secciones clave
# ═══════════════════════════════════════════════════════════════════════════════

# 4.1 Problemática — párrafo de contexto general
problematica_contexto_nuevo = (
    "Durante el desarrollo de las actividades laborales en la Escuela Profesional "
    "de Odontología de la Universidad Privada de Tacna, se identificaron múltiples "
    "deficiencias en la gestión administrativa y operativa del Comité de Mejora "
    "Continua (CMC) que afectaban significativamente la eficiencia del área. El "
    "personal de la oficina carecía de una herramienta digital centralizada para "
    "registrar, asignar y hacer seguimiento de las tareas diarias, lo que "
    "generaba pérdida de información, retrasos en la entrega de actividades y "
    "dificultades para elaborar reportes mensuales de productividad."
)

# 4.1 párrafo de estructura
problematica_estructura_nuevo = (
    "La oficina contaba con personal en roles diferenciados (Jefa y Asistente), "
    "actividades recurrentes de gestión institucional y la necesidad de generar "
    "evidencia documental periódica de las tareas realizadas. Esta estructura "
    "demandaba procesos digitalizados que garantizaran acceso rápido al estado "
    "de las tareas, facilitaran la toma de decisiones y permitieran generar "
    "reportes automáticos sin depender de hojas de cálculo manuales."
)

# 4.2.1 Problema General
problema_general_nuevo = (
    "¿De qué manera la ausencia de un sistema web centralizado de gestión de "
    "tareas afecta la eficiencia operativa, la trazabilidad de actividades y la "
    "capacidad de generación de reportes mensuales de productividad del personal "
    "del Comité de Mejora Continua de la Escuela Profesional de Odontología de "
    "la Universidad Privada de Tacna?"
)

# 4.4.1 Objetivo General
objetivo_general_nuevo = (
    "Desarrollar e implementar un Sistema Web de Gestión de Tareas para la "
    "Oficina del Comité de Mejora Continua (CMC) de la Escuela Profesional de "
    "Odontología de la Universidad Privada de Tacna, mediante arquitectura "
    "cliente-servidor con stack tecnológico MERN (MongoDB, Express.js, React y "
    "Node.js) y API RESTful, que centralice y automatice el registro, seguimiento "
    "y reporte de tareas del personal, mejorando significativamente la eficiencia "
    "operativa mediante digitalización del flujo de trabajo y generación "
    "automática de reportes mensuales en PDF."
)

reemplazos_cap4 = {
    # 4.1
    "Durante el desarrollo de las actividades laborales en la Escuela Pro- fesional de Odontología de la":
        problematica_contexto_nuevo,
    "Durante el desarrollo de las actividades laborales en la Escuela Profesional de Odontología de la":
        problematica_contexto_nuevo,
    "La escuela cuenta con una población aproximada de estudiantes de pregrado":
        problematica_estructura_nuevo,
    # Reemplazar los sub-problemas de Odontología por los del sistema de tareas
    "Gestión Descentralizada de Prospectos Académicos":
        "Gestión Manual y Descentralizada de Tareas Administrativas",
    "captación y seguimiento de estudiantes potenciales (leads) se realizaba de manera completamente manual":
        (
            "El registro y seguimiento de tareas del personal de la oficina se realizaba "
            "de manera completamente manual mediante notas físicas, correos electrónicos "
            "dispersos y hojas de cálculo independientes, sin un sistema unificado que "
            "permitiera centralizar el estado de cada actividad, la persona responsable, "
            "la prioridad y la fecha de entrega. Esta metodología generaba pérdida de "
            "tareas, duplicidad de esfuerzos, dificultad para medir la productividad del "
            "área y ausencia de trazabilidad en el trabajo realizado."
        ),
    "Ausencia de Plataforma de Comunicación Institucional":
        "Ausencia de Sistema de Generación de Reportes Automáticos",
    "No existía un sistema centralizado para publicación y difusión de noticias institucionales":
        (
            "No existía un mecanismo automatizado para generar reportes mensuales de "
            "las tareas completadas por el personal. La elaboración de estos informes "
            "exigía recopilar manualmente información dispersa, invertir horas en "
            "consolidar datos y formatear documentos, con alto riesgo de omisiones y "
            "errores. La ausencia de reportes sistemáticos dificultaba la evaluación "
            "de la productividad del área y la presentación de evidencias ante los "
            "organismos de acreditación institucional."
        ),
    "Control Manual de Inscripciones Estudiantiles":
        "Falta de Diferenciación de Roles y Control de Acceso",
    "El proceso de inscripción de nuevos estudiantes se ejecutaba mediante formularios físicos":
        (
            "No existía una plataforma que diferenciara los niveles de acceso y "
            "permisos según el rol del usuario (Jefa o Asistente). Toda la información "
            "era accesible de forma indiscriminada o dependía de credenciales compartidas "
            "de manera insegura. Esta carencia impedía implementar flujos de trabajo "
            "con responsabilidades definidas, como restringir la edición o eliminación "
            "de tareas únicamente al rol de Jefa, y dificultaba la auditoría de acciones "
            "realizadas por cada usuario."
        ),
    "Ausencia de Sistema de Gestión de Eventos Académicos":
        "Ausencia de Registro Histórico y Estadísticas de Productividad",
    "La programación, publicación y control de eventos académicos carecía de plataforma":
        (
            "No existía un repositorio digital que almacenara el historial completo "
            "de tareas por período, impidiendo consultar el trabajo realizado en meses "
            "anteriores y generar estadísticas de productividad comparativas. Esta "
            "carencia impedía identificar carga de trabajo por usuario, medir tendencias "
            "de cumplimiento y generar indicadores cuantitativos para la mejora continua "
            "del área."
        ),
    # 4.2.1 Problema General
    "¿De qué manera la ausencia de una plataforma web integral centralizada":
        problema_general_nuevo,
    # 4.4.1 Objetivo General
    "Desarrollar e implementar un Sistema Integral de Gestión Aca- démica (SIGA OdontologÝa) mediante arq":
        objetivo_general_nuevo,
    "Desarrollar e implementar un Sistema Integral de Gestión Académica (SIGA Odontología) mediante arq":
        objetivo_general_nuevo,
    # Nombre sistema en apertura Cap IV
    "Sistema Integral de Gestión Académica (SIGA Odontología), una aplicación web full-stack basada en ar":
        (
            "Sistema de Gestión de Tareas – Oficina EPO, una aplicación web full-stack "
            "basada en arquitectura cliente-servidor con stack tecnológico MERN "
            "(MongoDB, Express.js, React y Node.js), implementando API RESTful, sistema "
            "de autenticación y autorización mediante JWT, gestión de tareas con "
            "prioridades diferenciadas, generación automática de reportes PDF mensuales "
            "y diseño responsive, garantizando escalabilidad, seguridad y eficiencia en "
            "la digitalización de procesos administrativos del área."
        ),
}

found_cap4 = 0
for para in doc.paragraphs:
    txt = para.text
    for old_frag, new_text in reemplazos_cap4.items():
        key = old_frag[:40].lower()
        if key in txt.lower():
            set_para_text(para, new_text)
            found_cap4 += 1
            print(f"    ✓ Cap IV - Reemplazado: '{old_frag[:55]}'")
            break

print(f"  Capítulo IV: {found_cap4} reemplazo(s)")

# ═══════════════════════════════════════════════════════════════════════════════
# CAMBIO 6 — Marco Teórico: quitar tecnologías que no usa tu sistema
# y agregar las que sí usa
# ═══════════════════════════════════════════════════════════════════════════════

reemplazos_marco = {
    # React Hook Form + Zod → jsPDF + jsPDF-autotable
    "React Hook Form + Zod": "jsPDF + jsPDF-autotable",
    "React Hook Form junto con Zod": "jsPDF junto con jsPDF-autotable",
    "Gestión de Formularios: React Hook Form + Zod": "Generación de Reportes PDF: jsPDF + jsPDF-autotable",
    # Multer → Recharts
    "Upload de Archivos: Multer": "Visualización de Datos: Recharts",
    "Multer es middleware de Node.js para manejo de multipart/form-data": (
        "Recharts es una biblioteca de gráficos para React basada en D3.js que "
        "permite crear visualizaciones interactivas de datos como gráficos de "
        "barras, líneas y áreas. En el Sistema de Gestión de Tareas – Oficina EPO, "
        "Recharts se utiliza en el dashboard para visualizar estadísticas de "
        "productividad: total de tareas, tareas completadas, pendientes y de alta "
        "prioridad, facilitando la toma de decisiones basada en datos."
    ),
    # express-validator → date-fns / moment
    "Validación: express-validator": "Manejo de Fechas: date-fns y Moment.js",
    "express-validator es un conjunto de middlewares de Express basados en validator.js": (
        "date-fns y Moment.js son bibliotecas JavaScript para manejo, formateo y "
        "manipulación de fechas. En el Sistema de Gestión de Tareas – Oficina EPO, "
        "se utilizan para formatear la fecha y hora de completado de tareas, "
        "calcular rangos de fechas para los reportes mensuales y mostrar fechas "
        "en formato legible en la interfaz de usuario."
    ),
    # Concurrently → React Beautiful DnD
    "Control de Concurrencia: Concurrently": "Arrastrar y Soltar: React Beautiful DnD",
    "Es una utilidad CLI que permite ejecutar múltiples comandos npm scripts de manera concurrente": (
        "React Beautiful DnD es una biblioteca de arrastrar y soltar (drag and drop) "
        "para React, accesible y de alto rendimiento. En el Sistema de Gestión de "
        "Tareas – Oficina EPO, se utiliza para permitir al usuario reorganizar "
        "visualmente las tareas pendientes mediante arrastre, ofreciendo una "
        "experiencia de usuario intuitiva y moderna para la gestión de prioridades."
    ),
    # En SIGA Odontología referencias en marco teórico
    "En SIGA Odontología, el SRS delimita casos de uso (registro de estudiantes, gestión de eventos": (
        "En el Sistema de Gestión de Tareas – Oficina EPO, el SRS delimita casos "
        "de uso (creación de tareas, cambio de estado, generación de reportes, "
        "gestión de usuarios), reglas (validación de datos, roles diferenciados, "
        "registro automático de fecha de completado) y restricciones (autenticación "
        "JWT, autorización por roles), habilitando una matriz de trazabilidad hacia "
        "diseño, pruebas y despliegue."
    ),
    "En SIGA Odontología se incluyen: gestionar prospectos académicos (leads)": (
        "En el Sistema de Gestión de Tareas – Oficina EPO se incluyen: crear, "
        "editar y eliminar tareas con título, descripción y prioridad diferenciada "
        "(normal/alta); marcar tareas como completadas con registro automático de "
        "fecha y hora; filtrar tareas por estado (pendiente/completada); generar "
        "reportes mensuales en PDF de tareas completadas; visualizar estadísticas "
        "de productividad en dashboard; gestionar notificaciones internas; y "
        "administrar usuarios con roles diferenciados (Jefa y Asistente)."
    ),
    "En SIGA Odontología, la API REST expone más de 40 end- points": (
        "En el Sistema de Gestión de Tareas – Oficina EPO, la API REST expone "
        "endpoints organizados por recursos: autenticación (/api/auth), tareas "
        "(/api/tasks), reportes (/api/reports) y notificaciones (/api/notifications), "
        "facilitando la separación de responsabilidades entre frontend y backend."
    ),
    "En SIGA Odontología, dotenv gestiona: URI de conexión a MongoDB Atlas con credenciales": (
        "En el Sistema de Gestión de Tareas – Oficina EPO, dotenv gestiona: URI "
        "de conexión a MongoDB Atlas con credenciales, secreto JWT para firma de "
        "tokens, puerto del servidor, URL del cliente para CORS, y modo de entorno "
        "(development/production). El archivo .env se incluye en .gitignore para "
        "evitar exposición de secretos en control de versiones."
    ),
    "En SIGA Odontología, XP facilita validar primero los flujos de mayor valor (gestión de estudiantes": (
        "En el Sistema de Gestión de Tareas – Oficina EPO, XP facilita validar "
        "primero los flujos de mayor valor (gestión de tareas, reportes mensuales, "
        "autenticación por roles) y ajustar reglas de negocio conforme aparecen "
        "nuevos requerimientos; su énfasis en pruebas y refactorización mejora la "
        "mantenibilidad a largo plazo."
    ),
}

for old_frag, new_text in reemplazos_marco.items():
    for para in doc.paragraphs:
        if old_frag[:40].lower() in para.text.lower():
            set_para_text(para, new_text)
            print(f"    ✓ Marco Teórico - '{old_frag[:55]}'")
            break

# ═══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(DEST)
print("\n✅ INFORME.docx guardado correctamente con todos los cambios.")
print(f"   Backup disponible en: INFORME_BACKUP.docx")

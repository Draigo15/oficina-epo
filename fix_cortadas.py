# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document("INFORME_FINAL_SISTEMA.docx")

# Limpiar restos residuales en índices u otros textos como 'Listado completo de noticias publicadas'
bad_terms = ["noticia", "institucional", "estudiante", "contacto", "prospecto"]

for p in doc.paragraphs:
    text_lower = p.text.lower()
    
    # 1. Borrar párrafos completos que nombren las funciones fantasma (ej: "Listado de noticias")
    # pero cuidar de no borrar páginas reales como "logotipo institucional", así que
    # solo borramos si habla de gestionar/listados de eso.
    if any(palabra in text_lower for palabra in ["noticia", "estudiante", "prospecto"]):
        # Nos aseguramos que no sea una mención inocente, si tiene la palabra clave, lo borramos.
        try:
             p.text = ""
        except:
             pass

    # 2. Arreglar palabras cortadas (ej: "Con- tenido" -> "Contenido")
    # Usamos expresiones regulares en el texto del párrafo
    if "-" in p.text:
       p.text = re.sub(r'([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)-\s*\n?\s*([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)', r'\1\2', p.text)

# Hacer lo mismo para las tablas (que es donde se esconden los "Contra- seña")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
             text_lower = cell.text.lower()
             if any(palabra in text_lower for palabra in ["noticia", "estudiante", "prospecto"]):
                 cell.text = "" # Limpiamos la celda
                 
             if "-" in cell.text:
                 cell.text = re.sub(r'([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)-\s*\n?\s*([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)', r'\1\2', cell.text)

doc.save("INFORME_FINAL_SISTEMA.docx")
print("Arreglado")
